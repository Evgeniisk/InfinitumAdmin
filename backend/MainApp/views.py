from django.shortcuts import render
from django.http import JsonResponse, HttpResponse, HttpResponseBadRequest
import json
#The login_required decorator does the following:
#If the user isn't logged in, it redirects the user to settings.LOGIN_URL, passing the path specified by setting.Login_URL in the browser.
#If the user is logged in, it executes the view normally. The view code is then free to assume the user is logged in because the decorator does all the checking and security.
from django.contrib.auth.decorators import login_required
from .models import UploadedContractWordDocumentIndividual, UploadedContractWordDocumentCompany, UploadedInvoiceWordDocumentCompany, UploadedInvoiceWordDocumentIndividual, Client_Individual, Client_Company, Represents, Contract, Invoice_Summary, InvoiceItem, ContractItem, Job, JobAssigned, Invoice, InvoicesToInvoiceItems, EmailInvoiceWithoutContract_individual, EmailInvoiceFirstWithContract_individual, EmailInvoiceOtherInvoicesWithContract_individual, EmailInvoiceReminder_individual, EmailInvoiceWithoutContract_company, EmailInvoiceFirstWithContract_company, EmailInvoiceOtherInvoicesWithContract_company, EmailInvoiceReminder_company, EmailInvoiceOnlyOneWithContract_company, EmailInvoiceOnlyOneWithContract_individual, EmailSubjectforDocuSign_company, EmailSubjectforDocuSign_individual, ContactsCompaniesToBeTodThrough, ContactsIndividualsToBeTodThrough, ContactsCompaniesToBeCCdThrough, ContactsIndividualsToBeCCdThrough, ContactsCompaniesToBeBccdThrough, ContactsIndividualsToBeBccdThrough, UsersToBeTodThrough, UsersToBeCCdThrough, UsersToBeBccdThrough
from django.views.decorators.csrf import csrf_exempt
from django.contrib.contenttypes.models import ContentType
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from copy import deepcopy
from django.conf import settings
from urllib.parse import urlparse, parse_qs
import re
from AuthenticationApp.models import CustomUser
import tempfile
from django.core.files import File
from django.utils import timezone
from datetime import datetime, timedelta
import base64
import requests
from AuthenticationApp.models import DocuSignToken
from docx.shared import RGBColor
import subprocess
import os
from django.core.mail import EmailMessage
import time
from dateutil.relativedelta import relativedelta
from django.views.generic import TemplateView

#Entry point for production Vue template:
#class VueAppView(TemplateView):
#    template_name = 'MainApp/spa/index.html'


# Create your views here.
#Function for detecting and returning placeholders in company email templates:
def EmailCompanyPlaceholdersConfirmationFunction(String):
    valid_placeholders = {
        "{Company_Name}",
        "{Signer_First_Name}",
        "{Signer_Last_Name}",
        "{Period}",
        "{Related_Contact_First_Name}",
        "{Related_Contact_Last_Name}",
        "{Related_Contact_Company_Name}"
    }
    #Finds all placeholders inside the {} braces
    Detected_Placeholders = re.findall(r'\{.*?\}', String)
    #Returns only the placeholders that are valid
    return [placeholder for placeholder in Detected_Placeholders if placeholder in valid_placeholders]

#Function for detecting and returning placeholders in individual email templates:
def EmailIndividualPlaceholdersConfirmationFunction(String):
    valid_placeholders = {
        "{Client_First_Name}",
        "{Client_Last_Name}",
        "{Period}",
        "{Related_Contact_First_Name}",
        "{Related_Contact_Last_Name}",
        "{Related_Contact_Company_Name}"
    }
    #Finds all placeholders inside the {} braces
    Detected_Placeholders = re.findall(r'\{.*?\}', String)
    #Returns only the placeholders that are valid
    return [placeholder for placeholder in Detected_Placeholders if placeholder in valid_placeholders]

#Function to replace placeholders that are valid in the text for emails meant for companies and return them replaced:
def ReplacePlaceholdersEmailCompany(String, replacements):
    valid_placeholders = {
        "{Company_Name}",
        "{Signer_First_Name}",
        "{Signer_Last_Name}",
        "{Period}",
        "{Related_Contact_First_Name}",
        "{Related_Contact_Last_Name}",
        "{Related_Contact_Company_Name}"
    }
    #Finds all placeholders inside the {} braces
    Detected_Placeholders = re.findall(r'\{.*?\}', String)
    for placeholder in Detected_Placeholders:
        #Replaces only if the placeholder is in valid placeholders
        if placeholder in valid_placeholders:
            String = String.replace(placeholder, replacements.get(placeholder, placeholder))
    return String

#Function to replace placeholders that are valid in the text for emails meant for individuals and return them replaced:
def ReplacePlaceholdersEmailIndividual(String, replacements):
    valid_placeholders = {
        "{Client_First_Name}",
        "{Client_Last_Name}",
        "{Period}",
        "{Related_Contact_First_Name}",
        "{Related_Contact_Last_Name}",
        "{Related_Contact_Company_Name}"
    }
    #Finds all placeholders inside the {} braces
    Detected_Placeholders = re.findall(r'\{.*?\}', String)
    for placeholder in Detected_Placeholders:
        #Replaces only if the placeholder is in valid placeholders
        if placeholder in valid_placeholders:
            String = String.replace(placeholder, replacements.get(placeholder, placeholder))
    return String

#Functions to handle GET, POST, PUT and DELETE requests for email templates
#This function will be used to in email templates api endpoint functions for both company and individual templates to process requests
def EmailTemplateHandlingFunction(request, database_table):
    if request.method == 'GET':
        #Queries the database to return the email template for the workplace account that the user is associated with
        email_template = database_table.objects.get(workplace=request.user.workplace)
        return JsonResponse(email_template.as_dict())
    
    elif request.method == 'POST':
        #Retrieves the request data and turns it into python readable dictionary
        data = json.loads(request.body)
        subject = data.get('Subject')
        body = data.get('Body')
        #Created new template in the database
        email_template = database_table.objects.create(subject=subject, body=body, workplace = request.user.workplace)
        #Checks for which client type the template creation request is sent
        #Calls relevant confirmation functions depending on the client type the template is created for
        if database_table in [EmailInvoiceWithoutContract_company, EmailInvoiceFirstWithContract_company, EmailInvoiceOtherInvoicesWithContract_company, EmailInvoiceReminder_company, EmailInvoiceOnlyOneWithContract_company]:
            Confirmation_Placeholders_in_Subject = EmailCompanyPlaceholdersConfirmationFunction(subject)  #Calling the placeholders confirmation function.
            Confirmation_Placeholders_in_Body = EmailCompanyPlaceholdersConfirmationFunction(body)
        else:
            Confirmation_Placeholders_in_Subject = EmailIndividualPlaceholdersConfirmationFunction(subject)
            Confirmation_Placeholders_in_Body = EmailIndividualPlaceholdersConfirmationFunction(body)
        return JsonResponse({'email_template': email_template.as_dict(), 'Confirmation_Placeholders_in_Subject': Confirmation_Placeholders_in_Subject, 'Confirmation_Placeholders_in_Body': Confirmation_Placeholders_in_Body})
    
    elif request.method == 'PUT':
        #Retrieves the request data and turns it into python readable dictionary
        data = json.loads(request.body)
        subject = data.get('Subject')
        body = data.get('Body')
        #Modifies the existing template
        email_template = database_table.objects.get(workplace = request.user.workplace)
        email_template.subject = subject
        email_template.body = body
        email_template.save()
        #Checks for which client type the template creation request is sent
        #Calls relevant confirmation functions depending on the client type the template is created for
        if database_table in [EmailInvoiceWithoutContract_company, EmailInvoiceFirstWithContract_company, EmailInvoiceOtherInvoicesWithContract_company, EmailInvoiceReminder_company, EmailInvoiceOnlyOneWithContract_company]:
            Confirmation_Placeholders_in_Subject = EmailCompanyPlaceholdersConfirmationFunction(subject)
            Confirmation_Placeholders_in_Body = EmailCompanyPlaceholdersConfirmationFunction(body)
        else:
            Confirmation_Placeholders_in_Subject = EmailIndividualPlaceholdersConfirmationFunction(subject)
            Confirmation_Placeholders_in_Body = EmailIndividualPlaceholdersConfirmationFunction(body)
        return JsonResponse({'email_template': email_template.as_dict(), 'Confirmation_Placeholders_in_Subject': Confirmation_Placeholders_in_Subject, 'Confirmation_Placeholders_in_Body': Confirmation_Placeholders_in_Body})
    
    elif request.method == 'DELETE':
        #Queries the template and deletes it from the database
        email_template = database_table.objects.get(workplace=request.user.workplace)
        email_template.delete()
        return JsonResponse({'object': 'DELETED'})
    
    else:
        return HttpResponse(status=405)


#API endpoint which handles changes to Email Invoice Without Contract template using the EmailTemplateHandlingFunction
@csrf_exempt
@login_required
def api_EmailInvoiceWithoutContract(request):
    #gets the url header from the request
    url = request.headers.get('Url-Header')
    type = url.split("/")[-1] #extracts which client type the template is for from the url's last segment in its path
    #calls appropriate emailtemplatehandlingfunction based on the client type to process the request
    if type == 'Company':
        return EmailTemplateHandlingFunction(request, database_table=EmailInvoiceWithoutContract_company)
    else:
        return EmailTemplateHandlingFunction(request, database_table=EmailInvoiceWithoutContract_individual)

#API endpoint which handles changes to Email Invoice First With Contract template using the EmailTemplateHandlingFunction
@csrf_exempt
@login_required
def api_EmailInvoiceFirstWithContract(request):
    #gets the url header from the request
    url = request.headers.get('Url-Header')
    type = url.split("/")[-1] #extracts which client type the template is for from the url's last segment in its path
    #calls appropriate emailtemplatehandlingfunction based on the client type to process the request
    if type == 'Company':
        return EmailTemplateHandlingFunction(request, database_table=EmailInvoiceFirstWithContract_company)
    else:
        return EmailTemplateHandlingFunction(request, database_table=EmailInvoiceFirstWithContract_individual)

#API endpoint which handles changes to Email Invoice Other Invoices With Contract template using the EmailTemplateHandlingFunction
@csrf_exempt
@login_required
def api_EmailInvoiceOtherInvoicesWithContract(request):
    #gets the url header from the request
    url = request.headers.get('Url-Header')
    type = url.split("/")[-1] #extracts which client type the template is for from the url's last segment in its path
    #calls appropriate emailtemplatehandlingfunction based on the client type to process the request
    if type == 'Company':
        return EmailTemplateHandlingFunction(request, database_table=EmailInvoiceOtherInvoicesWithContract_company)
    else:
        return EmailTemplateHandlingFunction(request, database_table=EmailInvoiceOtherInvoicesWithContract_individual)

#API endpoint which handles changes to Email Invoice Reminder template using the EmailTemplateHandlingFunction
@csrf_exempt
@login_required
def api_EmailInvoiceReminder(request):
    #gets the url header from the request
    url = request.headers.get('Url-Header')
    type = url.split("/")[-1] #extracts which client type the template is for from the url's last segment in its path
    #calls appropriate emailtemplatehandlingfunction based on the client type to process the request
    if type == 'Company':
        return EmailTemplateHandlingFunction(request, database_table=EmailInvoiceReminder_company)
    else:
        return EmailTemplateHandlingFunction(request, database_table=EmailInvoiceReminder_individual)

#API endpoint which handles changes to Email Invoice Only One With Contract template using the EmailTemplateHandlingFunction
@csrf_exempt
@login_required
def api_EmailInvoiceOnlyOneWithContract(request):
    #gets the url header from the request
    url = request.headers.get('Url-Header')
    type = url.split("/")[-1] #extracts which client type the template is for from the url's last segment in its path
    #calls appropriate emailtemplatehandlingfunction based on the client type to process the request
    if type == 'Company':
        return EmailTemplateHandlingFunction(request, database_table=EmailInvoiceOnlyOneWithContract_company)
    else:
        return EmailTemplateHandlingFunction(request, database_table=EmailInvoiceOnlyOneWithContract_individual)
    

#API endpoint for handling requests to manage Email Subject for Docusign templates
@csrf_exempt
@login_required
def api_EmailSubjectforDocuSign(request):
    #gets the url header from the request
    url = request.headers.get('Url-Header')
    type = url.split("/")[-1] #extracts which client type the template is for from the url's last segment in its path
    
    #Processes the requests according to the method and which client type the template is for based on the last segment in the extracted url's path
    if request.method == 'GET':
        if type == 'Company':
            email_template = EmailSubjectforDocuSign_company.objects.get(workplace=request.user.workplace)
            return JsonResponse(email_template.as_dict())
        else:
            email_template = EmailSubjectforDocuSign_individual.objects.get(workplace=request.user.workplace)
            return JsonResponse(email_template.as_dict())
    elif request.method == 'POST':
        #Retrieves the request data and turns it into python readable dictionary
        data = json.loads(request.body)
        if type == 'Company':
            email_template = EmailSubjectforDocuSign_company.objects.create(subject = data, workplace=request.user.workplace)
            Confirmation_Placeholders_in_Subject = EmailCompanyPlaceholdersConfirmationFunction(email_template.subject)
            return JsonResponse({'email_template': email_template.as_dict(), 'Confirmation_Placeholders_in_Subject': Confirmation_Placeholders_in_Subject})
        else:
            email_template = EmailSubjectforDocuSign_individual.objects.create(subject = data, workplace=request.user.workplace)
            Confirmation_Placeholders_in_Subject = EmailIndividualPlaceholdersConfirmationFunction(email_template.subject)
            return JsonResponse({'email_template': email_template.as_dict(), 'Confirmation_Placeholders_in_Subject': Confirmation_Placeholders_in_Subject})
    elif request.method == 'PUT':
        #Retrieves the request data and turns it into python readable dictionary
        data = json.loads(request.body)
        if type == 'Company':
            email_template = EmailSubjectforDocuSign_company.objects.get(workplace=request.user.workplace)
            email_template.subject = data
            email_template.save()
            Confirmation_Placeholders_in_Subject = EmailCompanyPlaceholdersConfirmationFunction(email_template.subject)
            return JsonResponse({'email_template': email_template.as_dict(), 'Confirmation_Placeholders_in_Subject': Confirmation_Placeholders_in_Subject})
        else:
            email_template = EmailSubjectforDocuSign_individual.objects.get(workplace=request.user.workplace)
            email_template.subject = data
            email_template.save()
            Confirmation_Placeholders_in_Subject = EmailIndividualPlaceholdersConfirmationFunction(email_template.subject)
            return JsonResponse({'email_template': email_template.as_dict(), 'Confirmation_Placeholders_in_Subject': Confirmation_Placeholders_in_Subject})
        return JsonResponse(email_template.as_dict())
    elif request.method == 'DELETE':
        if type == 'Company':
            email_template = EmailSubjectforDocuSign_company.objects.get(workplace=request.user.workplace)
        else:
            email_template = EmailSubjectforDocuSign_individual.objects.get(workplace=request.user.workplace)
        email_template.delete()
        return JsonResponse({'Object': 'DELETED'})


    


#API endpoint for handling requests to add clients in a table on the front end, see them from a table on the front end, and delete them from the database (both company and individuals)
@csrf_exempt
@login_required
def api_clients(request):
    if request.method == 'GET':
        #Retrieves all individual client from the database and puts them into a list
        individual_clients = [
            {**client.as_dict(), 'ClientType': 'Client_Individual'}
            for client in Client_Individual.objects.filter(workplace=request.user.workplace)
        ]
        #Retrieves all company clients from the database and puts them into a list
        company_clients = [
            {**client.as_dict(), 'ClientType': 'Client_company'}
            for client in Client_Company.objects.filter(workplace=request.user.workplace)
        ]
        #Combines the two lists into one list
        all_clients = individual_clients + company_clients
        return JsonResponse({'clients': all_clients})
    
    elif request.method == 'POST':
        #Retrieves the request data and turns it into python readable dictionary
        data = json.loads(request.body)
        #Retrieved a ClientType value from the dictionary to process based on this value
        client_type = data.get('ClientType')
        if client_type == 'Client_Company':
            #Extracts company details
            name = data.get('CompanyName')
            email = data.get('CompanyEmail')
            phone = data.get('Phone')
            address = data.get('address')
            individual = data.get('IndividualId')
            #If individual id has some data in the request dictionary (the id) the individual client object is retrieved from the databased and assigned to the signer attribue via foreign key.
            #This signer field is specific to company client only because a company has to have an individual who can sign contracts on behalf of the company.
            if individual != '':
                individual = Client_Individual.objects.get(id=int(individual), workplace=request.user.workplace)
                company = Client_Company.objects.create(name=name, email=email, phone=phone, address=address, Signer=individual, workplace=request.user.workplace)
            else:
                company = Client_Company.objects.create(name=name, email=email, phone=phone, address=address, workplace=request.user.workplace)
            return JsonResponse(company.as_dict())
    
        elif client_type == 'Client_Individual':
            #Extracts individual's details
            Fname = data.get('Fname')
            Lname = data.get('Lname')
            phone = data.get('Phone')
            email = data.get('ClientEmail')
            address = data.get('address')
            workplace = request.user.workplace
            client_individual_post = Client_Individual.objects.create(Fname=Fname, Lname=Lname, phone=phone, address=address, email=email, workplace=workplace)  #Creates the client individual object in the database
            return JsonResponse(client_individual_post.as_dict())  #Returns this client for validation of client creation on the front end and dynamic entry to the client table on the front end.
        else:
            return HttpResponse(status=405)
    
    elif request.method == 'DELETE':
        #Retrieves the request data and turns it into python readable dictionary
        data = json.loads(request.body)
        #Retrieved a ClientType value from the dictionary to process based on this value
        client_type = data.get('ClientType')
        id = data.get('id')
        #Deletes the client from the database as per the request
        if client_type == 'Client_company':
            company = Client_Company.objects.get(id=id, workplace=request.user.workplace)
            company.delete()
            return JsonResponse({'Company': 'DELETED'})
        elif client_type == 'Client_Individual':
            individual = Client_Individual.objects.get(id=id, workplace=request.user.workplace)
            individual.delete()

#API endpoint for getting client details and modiying them for a single client
@csrf_exempt
@login_required
def api_client(request):
    if request.method == 'GET':
        #Getting the header
        url = request.headers.get('Url-Header')
        parsed_url = urlparse(url) #splits the url into its components (the path, quey, fragment etc..) that I can access
        query_params = parse_qs(parsed_url.query) #converts the query string of the url into a dictionary with keys being single strings such as 'id' and values being lists of values
        client_id = query_params.get('id')[0] #extracting the client id from the url by indexing the value of the 'id' key as [0] because the values are stored in a list
        ClientType = query_params.get('ClientType')[0] #extracting the client type from the url by indexing the value of the 'ClientType' key as [0] because the values are stored in a list
        print(ClientType)
        print(client_id)
        print(url)
        #Retrieving the requested client details from the database and returning them to the front end
        if ClientType == 'Client_Individual':
            Client = Client_Individual.objects.get(id=client_id, workplace=request.user.workplace)
            return JsonResponse(Client.as_dict())
        else:
            Client = Client_Company.objects.get(id=client_id, workplace=request.user.workplace)
            return JsonResponse(Client.as_dict())
    elif request.method == 'PUT':
        #Getting the header
        url = request.headers.get('Url-Header')
        parsed_url = urlparse(url) #splits the url into its components (the path, quey, fragment etc..) that I can access
        query_params = parse_qs(parsed_url.query) #converts the query string of the url into a dictionary with keys being single strings such as 'id' and values being lists of values
        client_id = query_params.get('id')[0] #extracting the client id from the url by indexing the value of the 'id' key as [0] because the values are stored in a list
        ClientType = query_params.get('ClientType')[0] #extracting the client type from the url by indexing the value of the 'ClientType' key as [0] because the values are stored in a list
        #Retrieves the request data and turns it into python readable dictionary
        data = json.loads(request.body)
        #Updates the client with the data send in the body of the request
        if ClientType == 'Client_Individual':
            Client = Client_Individual.objects.get(id=client_id, workplace=request.user.workplace)
            Fname = data.get('Fname')
            Lname = data.get('Lname')
            email = data.get('ClientEmail')
            phone = data.get('Phone')
            address = data.get('address')
            Client.Fname = Fname
            Client.Lname = Lname
            Client.email = email
            Client.phone = phone
            Client.address = address
            Client.save()
        else:
            Client = Client_Company.objects.get(id=client_id, workplace = request.user.workplace)
            name = data.get('CompanyName')
            email = data.get('CompanyEmail')
            phone = data.get('Phone')
            address = data.get('address')
            individual = data.get('IndividualId')
            if individual != '':
                individual = Client_Individual.objects.get(id=int(individual), workplace=request.user.workplace)
                Client.Signer = individual
            elif individual == '':
                try:
                    individual = Client.Signer
                    Client.Signer = None
                except:
                    pass
            Client.name = name
            Client.email = email
            Client.phone = phone
            Client.address = address
            Client.save()
        return JsonResponse(Client.as_dict())
    
#API endpoint to getting, creating and modifying represents relationships between clients (see MainApp.models file for reference) (see the frontend Client.vue file for reference)
@csrf_exempt
@login_required
def api_represents_now(request):
    if request.method == 'GET':
        #Getting the header
        url = request.headers.get('Url-Header')
        parsed_url = urlparse(url) #splits the url into its components (the path, quey, fragment etc..) that I can access
        query_params = parse_qs(parsed_url.query) #converts the query string of the url into a dictionary with keys being single strings such as 'id' and values being lists of values
        client_id = query_params.get('id')[0] #extracting the client id from the url by indexing the value of the 'id' key as [0] because the values are stored in a list
        client_type = query_params.get('ClientType')[0] #extracting the client type from the url by indexing the value of the 'ClientType' key as [0] because the values are stored in a list
        #Getting the content type for the requested client type
        from_content_type = ContentType.objects.get(model=client_type.lower())
        representslist = []
        #Retrives all Represents objects for the requested client
        represents = Represents.objects.filter(from_content_type=from_content_type, from_object_id=client_id, workplace=request.user.workplace)
        for item in represents:
            representslist.append(item.as_dict())

        return JsonResponse({"representslist": representslist})
    
    elif request.method == 'POST':
        #Retrieves the request data and turns it into python readable dictionary
        data = json.loads(request.body)
        #Extracts the from_type and from_id values from the request body dictionary
        from_type = data.get('from_type')
        from_id = data.get('from_id')
        #Gets the Django content type and model for the requested client that the relationships are being created for
        from_content_type = ContentType.objects.get(model=from_type.lower())
        from_model = from_content_type.model_class() #converts the contentype object into the django model class
        from_instance = from_model.objects.get(id=from_id)
        #Extracts the clients that were selected from the database to represent this client
        selected_contacts = data.get('selected_contacts')
        representsreturn = []
        #Loops through each client that was selected from the database to represent this client and creates a new represents relationship for this client with each selected client
        for contact in selected_contacts:
            to_type = contact['to_type']
            to_id = contact['to_id']
            to_content_type = ContentType.objects.get(model=to_type.lower())
            to_model = to_content_type.model_class() #converts the contentype object into the django model class
            to_instance = to_model.objects.get(id=to_id)

            represents = Represents.objects.create(
                from_content_type = from_content_type,
                from_object_id = from_id,
                to_content_type = to_content_type,
                to_object_id = to_id,
                workplace = request.user.workplace
            )
            representsreturn.append(represents.as_dict())
        return JsonResponse({'Represents': representsreturn})
    
    elif request.method == 'PUT':
        #Retrieves the request data and turns it into python readable dictionary
        data = json.loads(request.body)
        #Extracts the from_type and from_id values from the request body dictionary
        from_type = data.get('from_type')
        from_id = data.get('from_id')
        #Gets the Django content type and model for the requested client that the relationships are being modified for
        from_content_type = ContentType.objects.get(model=from_type.lower())
        from_model = from_content_type.model_class() #converts the contentype object into the django model class
        from_instance = from_model.objects.get(id=from_id)
        #Extracts the clients that were selected from the database to represent this client
        selected_contacts = data.get('selected_contacts')
        if selected_contacts:
            try:
                #Creates a set of client type and client id pares of the clients that the user chose to represent the client that the changes are being made for
                selected_contacts_set = set((contact['to_type'].lower(), contact['to_id']) for contact in selected_contacts)
                #Retrieves the existing represents entries of the client that the changes are being made for and stores them in a list
                represents = Represents.objects.filter(from_content_type=from_content_type, from_object_id=from_id, workplace=request.user.workplace)
                #Creates a set of client type and client id pares of the clients that currently (as in in the database before the changes made to it) represent the client that the changes are being made for
                represents_set = set((rep.to_content_type.model, rep.to_object_id) for rep in represents)
                #creates a set of clients that are not in the selected clients as requested by the user
                delete_set = represents_set - selected_contacts_set
                #Deletes the relationships that weren't selected by the client
                for model_name, to_id in delete_set:
                    to_content_type = ContentType.objects.get(model=model_name)
                    Represents.objects.filter(from_content_type=from_content_type, from_object_id=from_id, to_content_type=to_content_type, to_object_id=to_id, workplace=request.user.workplace).delete()
            except:
                pass

            representsreturn = []
            #Loops through the contacts selected by the client
            for contact in selected_contacts:
                to_type = contact['to_type'] #gets the client type of the selected client that was chosen to represent the client
                to_id = contact['to_id'] #gets the id of the selected client that was chosen to represent the client
                #Retrieves the content type of the selected client that was chosen to represent the client
                to_content_type = ContentType.objects.get(model=to_type.lower())
                to_model = to_content_type.model_class() #Retrieves the Django model class of the selected client's content type that was chosen to represent the client
                to_instance = to_model.objects.get(id=to_id) #Retreives the object instance of the selected client that was chosen to represent the client
                try:
                    #Tries to get the existing Represents entry for the client pairs
                    represents = Represents.objects.get(from_content_type=from_content_type, from_object_id=from_id, to_content_type=to_content_type, to_object_id=to_id, workplace=request.user.workplace)
                    #Updates the selected attribute of the represents entry if the database entry of it is different from the client request
                    new_selected = contact.get("selected", False)
                    if represents.selected != new_selected:
                        represents.selected = new_selected
                        represents.save()
                    representsreturn.append(represents.as_dict())
                except Represents.DoesNotExist:
                    #If there is no existing Represents entry for the client pairs (e.g. the try returns a DoesNotExist error) a new entry in the database is created for that relationship
                    represents_create = Represents.objects.create(
                        from_content_type = from_content_type,
                        from_object_id = from_id,
                        to_content_type = to_content_type,
                        to_object_id = to_id,
                        workplace = request.user.workplace,
                        selected = contact.get("selected", False)
                        )
                    representsreturn.append(represents_create.as_dict())
            return JsonResponse({'Represents': representsreturn})
        else:
            #If the selected contacts list is empty it deletes any represents relationship entries this client has from the database
            Represents.objects.filter(from_content_type=from_content_type, from_object_id=from_id, workplace=request.user.workplace).delete()
            return JsonResponse({'Objects': 'DELETED'})
        

@csrf_exempt
@login_required
def api_represents(request):
    if request.method == 'GET':
        return JsonResponse({
            'RepresentsList': [
                represents.as_dict() for represents in Represents.objects.filter(workplace = request.user.workplace)
            ]
        })

#A function for extracting placeholders in {} braces from a word document
def extract_placeholders(doc_path):
    doc = Document(doc_path) #Loads the word document and stores it in the doc object
    placeholder_regex = r"\{(.*?)\}" #Regular expression which matches any test inside the {} braces
    placeholders = [] #Initialised list to store all found placeholders

    #Searches for placeholers in the document paragraphs using the regular expression and append them to the placeholders list to store them
    for para in doc.paragraphs:
        matches = re.findall(placeholder_regex, para.text)
        placeholders.append(matches)
    
    #Searches for placeholders inside the tables of the word document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(placeholder_regex, cell.text)
                placeholders.append(matches)
    
    return placeholders #Returns the list of placeholders found

#A function which extracts placeholders from a given string
def extract_placeholders_from_text(text):
    return re.findall(r"{(.*?)}", text) #Returns a list of any placeholders which were located inside {} in the string

#A function which extracts all placeholders (any text entries found inside the {} brackets) from a Word (.docx) document.
def extract_placeholders_from_docx(doc_path):
    doc = Document(doc_path) #Loads the word document and stores it in the doc object
    placeholder_regex = r"\{(.*?)\}" #Regex which matches any test inside the {} braces
    placeholders  = [] #Initialised list to store the placeholders gound by the regex

    #Loops through paragraphs in the word document and searches for placeholders as specified above
    for para in doc.paragraphs:
        matches = re.findall(placeholder_regex, para.text)
        placeholders.extend(matches)
    
    #Loops through all tables in the word document and seaches for placeholders inside table's cell's paragraphs as specified above
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = extract_placeholders_from_text(para.text)
                    placeholders.extend(matches)

    return placeholders #Returns the list of placeholders







#API endpoint for handling client requests with uploading contract word document templates
@csrf_exempt
@login_required
def api_upload_contract_worddocument(request):
    #Getting the header
    url = request.headers.get('Url-Header')
    parsed_url = urlparse(url) #splits the url into its components (the path, quey, fragment etc..) that I can access
    query_params = parse_qs(parsed_url.query) #converts the query string of the url into a dictionary with keys being single strings such as 'ClientType' and values being lists of values
    client_type = query_params.get('ClientType', [None])[0] #extracting the client type from the url by indexing the value of the 'ClientType' key as [0] because the values are stored in a list
    #Getting the header again because the previously retrieved header was modified
    url = request.headers.get('Url-Header')
    parsed_url = urlparse(url) #splits the url into its components (the path, quey, fragment etc..) that I can access
    #Extracting the path segments from the url
    path_segments = parsed_url.path.strip("/").split("/")
    #Getting the client type from the url by indexing the last path segment
    type = path_segments[-1]
    print(type)
    if request.method == 'GET':
        #Returns the file name of the uploaded document based on client type it is used for
        if client_type == "Client_company":
            file = UploadedContractWordDocumentCompany.objects.get(workplace=request.user.workplace)
            file_name = file.file.name.split("/")[-1]
            return JsonResponse({'FileName': file_name})
        elif client_type == "Client_Individual":
            file = UploadedContractWordDocumentIndividual.objects.get(workplace=request.user.workplace)
            file_name = file.file.name.split("/")[-1]
            return JsonResponse({'FileName': file_name})
        elif type == 'Company':
            file = UploadedContractWordDocumentCompany.objects.get(workplace=request.user.workplace)
            file_name = file.file.name.split("/")[-1]
            return JsonResponse({'FileName': file_name})
        else:
            file = UploadedContractWordDocumentIndividual.objects.get(workplace=request.user.workplace)
            file_name = file.file.name.split("/")[-1]
            return JsonResponse({'FileName': file_name})
    elif request.method == 'POST':
        #Retries the uploaded file from the request from data using the file key (refer to frontend/src/pages/TemplatesCompany.vue)
        doc_file = request.FILES["file"]
        #Extracts placeholders from the uploaded file
        placeholders = extract_placeholders_from_docx(doc_file)
        #Deletes the existing file in the database if it exists and saves the uploaded file
        if type == 'Company':
            try:
                current_file = UploadedContractWordDocumentCompany.objects.get(workplace = request.user.workplace)
                current_file.file.delete()
                current_file.delete()
            except UploadedContractWordDocumentCompany.DoesNotExist:
                pass
            doc_instance = UploadedContractWordDocumentCompany.objects.create(workplace = request.user.workplace, file = doc_file)
            file_name = doc_instance.file.name.split("/")[-1]
        else:
            try:
                current_file = UploadedContractWordDocumentIndividual.objects.get(workplace = request.user.workplace)
                current_file.file.delete()
                current_file.delete()
            except UploadedContractWordDocumentIndividual.DoesNotExist:
                pass
            doc_instance = UploadedContractWordDocumentIndividual.objects.create(workplace = request.user.workplace, file = doc_file)
            file_name = doc_instance.file.name.split("/")[-1]
        return JsonResponse({'FileName': file_name, 'Placeholders': placeholders}) #Returns the uploaded file name and detected placeholders inside it

#API endpoint for handling client requests with uploading invoice word document templates
@csrf_exempt
@login_required
def api_upload_invoice_worddocument(request):
    #Getting the header
    url = request.headers.get('Url-Header')
    parsed_url = urlparse(url) #splits the url into its components (the path, quey, fragment etc..) that I can access
    query_params = parse_qs(parsed_url.query) #converts the query string of the url into a dictionary with keys being single strings such as 'ClientType' and values being lists of values
    client_type = query_params.get('ClientType', [None])[0] #extracting the client type from the url by indexing the value of the 'ClientType' key as [0] because the values are stored in a list
    #Getting the header again because the previously retrieved header was modified
    url = request.headers.get('Url-Header')
    parsed_url = urlparse(url) #splits the url into its components (the path, quey, fragment etc..) that I can access
    path_segments = parsed_url.path.strip("/").split("/") #Extracting the path segments from the url
    #Getting the client type from the url by indexing the last path segment
    type = path_segments[-1]
    print(type)
    if request.method == 'GET':
        #Returns the file name of the uploaded document based on client type it is used for
        if client_type == 'Client_company':
            file = UploadedInvoiceWordDocumentCompany.objects.get(workplace=request.user.workplace)
            file_name = file.file.name.split("/")[-1]
            return JsonResponse({'FileName': file_name})
        elif client_type == 'Client_Individual':
            file = UploadedInvoiceWordDocumentIndividual.objects.get(workplace=request.user.workplace)
            file_name = file.file.name.split("/")[-1]
            return JsonResponse({'FileName': file_name})
        elif type == 'Company':
            file = UploadedInvoiceWordDocumentCompany.objects.get(workplace=request.user.workplace)
            file_name = file.file.name.split("/")[-1]
            return JsonResponse({'FileName': file_name})
        else:
            file = UploadedInvoiceWordDocumentIndividual.objects.get(workplace=request.user.workplace)
            file_name = file.file.name.split("/")[-1]
            return JsonResponse({'FileName': file_name})
    elif request.method == 'POST':
        #Retries the uploaded file from the request from data using the file key (refer to frontend/src/pages/TemplatesCompany.vue)
        doc_file = request.FILES['file']
        #Extracts placeholders from the uploaded file
        placeholders = extract_placeholders_from_docx(doc_file)
        #Deletes the existing file in the database if it exists and saves the uploaded file
        if type == 'Company':
            try:
                current_file = UploadedInvoiceWordDocumentCompany.objects.get(workplace=request.user.workplace)
                current_file.file.delete()
                current_file.delete()
            except UploadedInvoiceWordDocumentCompany.DoesNotExist:
                pass
            doc_instance = UploadedInvoiceWordDocumentCompany.objects.create(workplace=request.user.workplace, file = doc_file)
            file_name = doc_instance.file.name.split("/")[-1]
        else:
            try:
                current_file = UploadedInvoiceWordDocumentIndividual.objects.get(workplace=request.user.workplace)
                current_file.file.delete()
                current_file.delete()
            except UploadedInvoiceWordDocumentIndividual.DoesNotExist:
                pass
            doc_instance = UploadedInvoiceWordDocumentIndividual.objects.create(workplace=request.user.workplace, file=doc_file)
            file_name = doc_instance.file.name.split("/")[-1]
        return JsonResponse({'FileName': file_name, 'Placeholders': placeholders}) #Returns the uploaded file name and detected placeholders inside it

#This function replaces an existing key in a dictionary with a new key both of which it receives as inputs when called
def replace_key(d, old_key, new_key):
    if old_key in d:
        d[new_key] = d.pop(old_key)
    return d #Returns a modified dictionary with the new key

#This function replaces placeholders in a Word document template with values which are passed to it
def replace_placeholders(input_file, replacements, table_entries, contract, total_table_entries):
    doc = Document(input_file.file.path) #Loads the word file which is going to be edited
    placeholder_regex = r"\{.*?\}" #Regular expression to match any text inside the {} placeholders

    #This function replaces placeholders in a paragraph
    def replace_in_paragraph(paragraph, replace_dict=None):
        #This assigns the passed replacements dictionary to be used as replacements if it is passed, otherwise it uses replacements passed in the parent function.
        current_replacements = replace_dict if replace_dict else replacements
        full_text = paragraph.text
        #This checks if all the placeholders specified below are present
        if (
            "{Total_price}" in full_text and
            "{Total_Price_with_VAT}" in full_text and
            "{Number}" in full_text and
            "{From_month_year}" in full_text and
            "{To_month_year}" in full_text and
            "{Period}" in full_text and
            "{Total_price_for_each_period}" in full_text and
            "{Total_Price_with_VAT_for_each_period}" in full_text
        ):
            #If this if statement evaluates to true, meaning all placeholders are present, 
            if str(current_replacements.get("{Number}", "")) == "1":
                full_text = (
                    "The amount of {Total_price} plus VAT = {Total_Price_with_VAT} "
                    "will be paid in {Number} instalment and will be payable on {From_month_year}."
                )
        #This replaces the placeholders in the paragraph using the replacements dictionary passed to this function is any placeholders are matched by the regular expression
        if re.search(placeholder_regex, full_text):
            #This loops over all keys (placeholders) and their values to replace them in the paragraph
            for placeholder, value in current_replacements.items():
                full_text = full_text.replace(str(placeholder), str(value))
            #This clears all existing text runs in the paragraph
            for run in paragraph.runs:
                run.text = ""
            #If the passed paragraph has runs, this replaces the original text with placeholders with the text which has palceholders replaced 
            if paragraph.runs:
                paragraph.runs[0].text = full_text
                run=paragraph.runs[0]
            else:
                #If no runs exist this adds a new run with the replaces text
                paragraph.add_run(full_text)
            #These are placeholders which DocuSign recognises which it can replace with signature and date fiels inside a document automatically
            docusign_placeholders = ["/sn_director/", "/ds_director/", "/sn_customer/", "/ds_customer/"]
            #If any of the DocuSign specific placeholders are in the replaced textI change their font color to white to hide them in the document because by default DocuSign will detect them, and if they are in white they will be invisible in the document.
            if any(ph in full_text for ph in docusign_placeholders):
                font = run.font
                font.color.rgb = RGBColor(255, 255, 255)
    #Here I loop over all paragraphs in the document and replace placeholders in them using the function above  
    for para in doc.paragraphs:
        replace_in_paragraph(para)
    #Here I loop over all tables in the document
    for table in doc.tables:
        #In every table I iterate over each row and each cell in each row and each paragraph in each cell
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    #If any placeholder from the passed dictionaty is in the paragraph I replace it in the paragraph using the the replacements dictionary
                    if any(str(placeholder) in para.text for placeholder in replacements.keys()):
                        replace_in_paragraph(para, replace_dict=replacements)
    #Here I loop over through all tables to insert relevant values replacement placeholder fiels in them
    for table in doc.tables:
        #Here I initialise a template row variable to copy a template row into it 
        template_row = None
        #Here I iterate over all rows in the table after the first row (I assume it is the header)
        for row in table.rows[1:]:
            found = False
            #I iterate over each cell in each row and each paragraph in each cell
            for cell in row.cells:
                for para in cell.paragraphs:
                    #I iterate over each dictionary key and value pair in the table entries dictionary passed to this function
                    for entry in table_entries:
                        #I replace the original keys with placeholders in the table_entries dictionary passed
                        replace_key(entry, 'Name', '{Service_or_Product_Name}')
                        replace_key(entry, 'Total_price', '{Price_with_VAT}')
                        replace_key(entry, 'Price', '{Price}')
                        replace_key(entry, 'VAT', '{VAT}')
                        #Here I check if entry is a dictionary and if any placeholders match the paragraph text
                        if isinstance(entry, dict):
                            if any(str(placeholder) in para.text for placeholder in entry.keys()):
                                #Here I mark this row as the template row to insert in the table
                                template_row = row
                                found = True
                                break
                    if found:
                        break
                if found:
                    break
            if found:
                    break
        #If a template row was found in the table, I insert new row for each entry that was passed in the table_entries dictionary
        if template_row:
            #Here I get the index position of the template row within the table structure
            idx = table._tbl.index(template_row._tr)
            #For each entry in the table entries list I
            for i, entry in enumerate(table_entries):
                #Create a deep copy to not modify the original entry
                new_entry = deepcopy(entry)
                #Replace keys in the entry with the corresponding placeholders
                replace_key(entry, 'Name', '{Service_or_Product_Name}')
                replace_key(entry, 'Total_price', '{Price_with_VAT}')
                replace_key(entry, 'Price', '{Price}')
                replace_key(entry, 'VAT', '{VAT}')
                #Make a deep copy of the template row to insert as a new row
                new_row = deepcopy(template_row)
                #I iterate over each cell and paragraph in each cell in the new row, replace placeholders with actual entry values
                for cell in new_row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, replace_dict=entry)
                #Insert each new row below the previously inserted one
                table._tbl.insert(idx + 1 + i, new_row._tr)
            #Remove the original template row after instering all new rows
            table._tbl.remove(template_row._tr)
    #Get the last row of the table which I assume to contain summaries for total values of a quote
    last_row = table.rows[-1]
    #Replace keys in the total table entries dictionary with their placeholders
    replace_key(total_table_entries, 'Price', '{Total_price}')
    replace_key(total_table_entries, 'VAT', '{Total_VAT}')
    replace_key(total_table_entries, 'Price_with_VAT', '{Total_Price_with_VAT}')
    #I iterate over each cell and paragraph in the last row and replace placeholders with total values
    for cell in last_row.cells:
        for para in cell.paragraphs:
            replace_in_paragraph(para, replace_dict=total_table_entries)

    #Here I create a temporary file to save the modified word document
    tmp_file =  tempfile.NamedTemporaryFile(suffix=".docx", delete=True)
    #Here I store the temporary file's name and close the file
    tmp_file_name = tmp_file.name
    tmp_file.close()
    #Here I save the edited document to the temporary file path
    doc.save(tmp_file_name)
    
    #Here I open the temporary file for reading in binary mode
    with open(tmp_file_name, 'rb') as f:
        #Here I save the file content into the contracts file field, naming it with the contract ID
        contract.file.save(f"contract_{contract.id}.docx", File(f), save=True)
    
    import os
    #Here I remove the temporary fule from disk after saving to the relevant contracts file field.
    os.remove(tmp_file_name)

#This function regenerates a DocuSign signing URL for a particular contract.
@csrf_exempt
@login_required
def singing_later(request):
    if request.method == 'POST':
        #Retrieves the request data and turns it into python readable dictionary
        data = json.loads(request.body)
        #Extracts the 'contract_id' value from the request body dictionary
        contract_id = data.get('contract_id')
        #Here I retrieve the requested contract from the database based on its id
        contract = Contract.objects.get(id=contract_id, workplace=request.user.workplace)
        #Here I retrieve the DocuSign token object which contains DocuSign account details associated with the account that the request came from
        DocuSignAccount = DocuSignToken.objects.get(workplace=request.user.workplace)
        #Here I extract DocuSign account details
        access_token = DocuSignAccount.access_token
        account_id = DocuSignAccount.account_id
        director_email = DocuSignAccount.user_email
        director_name = DocuSignAccount.user_full_name
        #This is a URL for DocuSign API
        base_url = "https://demo.docusign.net/restapi/v2.1"
        #Here I retrieve the id of the client the contract is for
        client_id = (contract.client_company.id if contract.client_company else contract.client_individual.id)
        #Here I retrieve the type of the client the contract if for
        client_type = ("Client_company" if contract.client_company else "Client_individual")

        #Here I build the request payload for the 
        recipient_view_request = {
            "userName": director_name, #Account holders full name
            "email": director_email, #Account holders email
            "recipientId": "1", #Unchanging recipient id which is required by DocuSign (always 1)
            "clientUserId": "12345", #Random user ID (Required by DocuSign)
            "authenticationMethod": "none", #Takes this in as standard authentication method
            #This is the redirect url after signing it completed
            "returnUrl": f"http://localhost:5173/Client/contracts?ClientType={client_type}&id={client_id}"
        }
        #Here I build the DocuSign API endpoint to ask docusign to generate a URL to sign the document for the account holder
        view_url = f"{base_url}/accounts/{account_id}/envelopes/{contract.envelope_id}/views/recipient"
        #Here I send the reqyest to DocuSign api asking it to generate a new signing URL for this contract for the account specified and store the response (url) in the response variable
        response = requests.post(view_url, headers={
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }, json=recipient_view_request)

        if not response.ok:
            return JsonResponse({"error": "Failed to regenerate signing URL"}, status=500)
        #Here I extract the response url DocuSign responded me with
        signing_url = response.json()["url"]
        #Here I save the regenerated signing URL to the contract
        contract.signingUrl = signing_url
        contract.save()
        #Here I return the new URL to the front end
        return JsonResponse({"signing_url_director": signing_url})

#This function is used to handle client requests to get contracts information, create a contract and view it via docusign, sign a contract at a later stage via a put request, delete a contract and all associated automations with it.
@csrf_exempt
@login_required
def api_contract(request):
    if request.method == 'GET':
        #Getting the header
        if request.headers.get('Url-Header'):
            url = request.headers.get('Url-Header')
            parsed_url = urlparse(url) #splits the url into its components (the path, quey, fragment etc..) that I can access
            query_params = parse_qs(parsed_url.query) #converts the query string of the url into a dictionary with keys being single strings such as 'ClientType' and values being lists of values
            client_id = query_params.get('id')[0] #extracting the client id from the url by indexing the value of the 'id' key as [0] because the values are stored in a list
            client_type = query_params.get('ClientType')[0] #extracting the client type from the url by indexing the value of the 'ClientType' key as [0] because the values are stored in a list
            #Returning a list of serialised contracts for the client
            if client_type == 'Client_company':
                return JsonResponse({'contracts': [contract.as_dict() for contract in Contract.objects.filter(workplace=request.user.workplace, client_company__id=client_id)]})
            else:
                return JsonResponse({'contracts': [contract.as_dict() for contract in Contract.objects.filter(workplace=request.user.workplace, client_individual__id=client_id)]})
        else:
            return JsonResponse({
                'contracts': [
                    contract.as_dict() for contract in Contract.objects.filter(workplace=request.user.workplace)
                ]
            })
    
    if request.method == 'POST':
        #Getting the header
        url = request.headers.get('Url-Header')
        parsed_url = urlparse(url) #splits the url into its components (the path, quey, fragment etc..) that I can access
        query_params = parse_qs(parsed_url.query) #converts the query string of the url into a dictionary with keys being single strings such as 'ClientType' and values being lists of values
        print(url) #Here from when I was debugging

        client_type = query_params.get('ClientType', [None])[0] #extracting the client id from the url by indexing the value of the 'id' key as [0] because the values are stored in a list
        client_id = query_params.get('id', [None])[0] #extracting the client type from the url by indexing the value of the 'ClientType' key as [0] because the values are stored in a list
        #Retrieves the request body and turns it into python readable dictionary
        data = json.loads(request.body)
        original_entries = data.get("ProductsOrServices") #Extracts the ProductsOrServices values from the request body dictionary
        #Copies each entry in the original entried above from the request body so that the original_entries object is not modified
        table_entries = [entry.copy() for entry in original_entries]
        #Loops over the table entries and pops Deadline and Employees
        for entry in table_entries:
            entry.pop('Deadline', None)
            entry.pop('Employees', None)
        #Extracts the Total and NumberOfEmployees values from the request body dictionary
        Total = data.get("Total")
        People = data.get("NumberOfEmployees")
        #Extracts the Invoice_Schedule value from the request body dictionary
        invoices_data = data.get("Invoice_Schedule")
        #Invoice_Schedule value is a dictionary with some nested dectionaries so here the keys get indexed that I need to get their values
        number_of_invoices = invoices_data['number_of_invoices']
        invoice_schedule = invoices_data['invoice_schedule']
        invoice_duration = invoices_data['duration']
        Total_Price_per_invoice = invoices_data['totalPrice']
        Total_withVAT_per_invoice = invoices_data['totalWithVAT']
        last_invoice_total_price = invoices_data['lastInvoice']['totalPrice']
        last_invoice_Total_with_VAT = invoices_data['lastInvoice']['totalWithVAT']

            
        
        if client_type == 'Client_company':
            client_company = Client_Company.objects.get(workplace=request.user.workplace, id=client_id)
            Signer = client_company.Signer
            company_name = client_company.name
            company_address = client_company.address
            print(company_address)
            Signer_First_Name = Signer.Fname
            print(Signer_First_Name)
            Signer_Last_Name = Signer.Lname
            print(Signer_Last_Name)
            #A dictionary with placeholders as keys and what they should be replaced with as their values
            replacements = {
                "{Company_Name}": company_name,
                "{Company_Address}": company_address,
                "{Signer_First_Name}": Signer_First_Name,
                "{Signer_Last_Name}": Signer_Last_Name,
                "{Directors_Signature}": "/sn_director/",
                "{Date_Signed_Director}": "/ds_director/",
                "{Signer_Signature}": "/sn_customer/",
                "{Date_Signed_Signer}": "/ds_customer/",
                "{Total_price}": Total['Price'],
                "{Total_Price_with_VAT}": Total['Price_with_VAT'],
            }
            print(replacements)
            #Gets the contract template from the database to be used as the input file in replace_placeholders function
            input_file = UploadedContractWordDocumentCompany.objects.get(workplace=request.user.workplace)
            #Creates a contract object in the database
            contract = Contract.objects.create(workplace=request.user.workplace, client_company=client_company, Price=Total["Price"], VAT=Total["VAT"], total_amount=Total["Price_with_VAT"], People=People)
            #Builds different payments breakdowns strings and adds them to the replacements dictionary based on different invoice schedule conditions
            if number_of_invoices <= 1:
                replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be paid in 1 installment and will payable on {contract.created_at.strftime('%d-%m-%Y')}."})
            elif number_of_invoices > 1 and last_invoice_total_price:
                if invoice_duration == "for 1 year":
                    end_date = contract.created_at + relativedelta(years=1)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. Per instalment amount of the first {number_of_invoices - 1} instalments is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}. The last instalment amount is {last_invoice_total_price} plus VAT = {last_invoice_Total_with_VAT}."})
                elif invoice_duration == "for 1 week":
                    end_date = contract.created_at + timedelta(weeks=1)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. Per instalment amount of the first {number_of_invoices - 1} instalments is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}. The last instalment amount is {last_invoice_total_price} plus VAT = {last_invoice_Total_with_VAT}."})
                elif invoice_duration == "for 1 month":
                    end_date = contract.created_at + relativedelta(months=1)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. Per instalment amount of the first {number_of_invoices - 1} instalments is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}. The last instalment amount is {last_invoice_total_price} plus VAT = {last_invoice_Total_with_VAT}."})
                elif invoice_duration == "for 1 quarter":
                    end_date = contract.created_at + relativedelta(months=3)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. Per instalment amount of the first {number_of_invoices - 1} instalments is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}. The last instalment amount is {last_invoice_total_price} plus VAT = {last_invoice_Total_with_VAT}."})
                elif invoice_duration == "for 5 years":
                    end_date = contract.created_at + relativedelta(years=5)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. Per instalment amount of the first {number_of_invoices - 1} instalments is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}. The last instalment amount is {last_invoice_total_price} plus VAT = {last_invoice_Total_with_VAT}."})

            else:
                if invoice_duration == "for 1 year":
                    end_date = contract.created_at + relativedelta(years=1)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. {invoice_schedule} amount is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}."})
                elif invoice_duration == "for 1 week":
                    end_date = contract.created_at + timedelta(weeks=1)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. {invoice_schedule} amount is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}."})
                elif invoice_duration == "for 1 month":
                    end_date = contract.created_at + relativedelta(months=1)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. {invoice_schedule} amount is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}."})
                elif invoice_duration == "for 1 quarter":
                    end_date = contract.created_at + relativedelta(months=3)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. {invoice_schedule} amount is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}."})
                elif invoice_duration == "for 5 years":
                    end_date = contract.created_at + relativedelta(years=5)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. {invoice_schedule} amount is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}."})
            #Calls the replace_placeholders function to replace placeholders in the contract template to create ready personalised contract and save it in the database (reference the replace_placeholders function)
            replace_placeholders(input_file, replacements, table_entries, contract, total_table_entries=Total)

            print(client_company.Signer.email) #Debugging 

            try:
                DocuSign_subject = EmailSubjectforDocuSign_company.objects.get(workplace = request.user.workplace)
                #Gets the ContentType object for the client company
                content_type = ContentType.objects.get_for_model(client_company)
                #Gets all the Represents objects from the database where the client company is the object that the contacts are assigned to and where those contacts have a selected flag (True) because it means they are chosen for correspondence with the client by the user.
                represents_list = Represents.objects.filter(from_content_type=content_type, from_object_id = client_company.id, workplace=request.user.workplace, selected=True)
                #Loops over the Represents objects list
                for rep in represents_list:
                    #Checks what client types the contacts are and retrives relevant information needed depending on what client types they are
                    if isinstance(rep.to_entity, Client_Individual):
                        Related_Contact_First_Name = rep.to_entity.Fname
                        Related_Contact_Last_Name = rep.to_entity.Lname
                    elif isinstance(rep.to_entity, Client_Company):
                        Related_Contact_Company_Name = rep.to_entity.name
                #replacements dictionary used to replace placeholders with the values as shown below
                replacements = {
                    "{Company_Name}": company_name,
                    "{Signer_First_Name}": Signer_First_Name,
                    "{Signer_Last_Name}": Signer_Last_Name,
                    "{Related_Contact_First_Name}": Related_Contact_First_Name,
                    "{Related_Contact_Last_Name}": Related_Contact_Last_Name,
                    "{Related_Contact_Company_Name}": Related_Contact_Company_Name,
                }
                #replaces placeholders in the email template from the database by calling the ReplacePlaceholdersCompany function and using the dictionary above (refer to the function)
                email_subject = ReplacePlaceholdersEmailCompany(DocuSign_subject.subject, replacements)
            except:
                email_subject = "Contract for Signature"
            #Opens the created contract file in read binary mode
            with open(contract.file.path, 'rb') as file:
                #Reads the file content into memory as bytes
                doc_bytes = file.read()
                #Encodes the file into based64 format and converts it to a utf-8 string because this is the required format for DocuSign API
                doc_base64 = base64.b64encode(doc_bytes).decode('utf-8')
                #Gets the DocuSign account detials from the database for the workplace the request is sent from
                DocuSignAccount = DocuSignToken.objects.get(workplace=request.user.workplace)
                access_token = DocuSignAccount.access_token
                account_id = DocuSignAccount.account_id
                director_email = DocuSignAccount.user_email
                director_name = DocuSignAccount.user_full_name
                base_url = "https://demo.docusign.net/restapi/v2.1" #API URL for DocuSign 
                #Builds the envelope definition to send to DocuSign api
                #The details are taken from DocuSign REST API documentation
                envelope_definition = {
                    "emailSubject": email_subject, #Subject of the email that the contract is sent in by DocuSign
                    #Documents to be included in the envelope (the contract in our case)
                    "documents": [{
                        "documentBase64": doc_base64, #The contract encoded in base64
                        "name": f"Contract_{contract.id}.docx", #Name of the contract
                        "fileExtension": "docx", #The file type of the contract
                        "documentId": f"{contract.id}" #Unique id of the document to be sent (the contract)
                    }],
                    #People who will recieve the contract
                    "recipients": {
                        #People who will recieve the contract and will need to sign it
                        "signers": [
                            {
                                #First recipient: the account holder of the DocuSign account connected to my web app
                                "email": director_email,
                                "name": director_name,
                                "recipientId": "1", #Recipient id, has to be unique in this specific envelope only therefore I can just leave it at this
                                "routingOrder": "1", #The order for who gets the contract first, in this case it is the account holder of the DocuSign account connected to InfinitumAdmin
                                "clientUserId": "12345", #DocuSign requirement for embedded signing
                                #Signature and date placeholders for this recipient (the account holder). This is telling DocuSign that there are these placeholders in the document and telling it to search for them and place the fields over them
                                #Here I set all coordinated to 0 since I have this placeholder in the contract I am sending to the fields will be placed directly over it
                                "tabs": {
                                    "signHereTabs":[{
                                        "anchorString": "/sn_director/", #Placeholder for the DocuSign account holder signature
                                        "anchorUnits": "pixels",
                                        "anchorXOffset": "0",
                                        "anchorYOffset": "0"
                                    }],
                                    "dateSignedTabs":[{
                                        "anchorString": "/ds_director/", #Placeholder for the DocuSign account holder signed date
                                        "anchorUnits": "pixels",
                                        "anchorXOffset": "0",
                                        "anchorYOffset": "0"
                                    }]
                                }
                            },
                            {
                                #Second recipient:the client to whom the contract will be sent after the account holder signs it
                                "email": str(client_company.Signer.email), 
                                "name": f"{client_company.Signer.Fname} {client_company.Signer.Lname}",
                                "recipientId": "2", #Recipient id, has to be unique in this specific envelope only therefore I can just leave it at this
                                "routingOrder": "2", #The order for who gets the contract second, in this case it is the client who the contract is intended for by InfinitumAdmin user
                                #Signature and date placeholders for this recipient (the account holder). This is telling DocuSign that there are these placeholders in the document and telling it to search for them and place the fields over them
                                #Here I set all coordinated to 0 since I have this placeholder in the contract I am sending to the fields will be placed directly over it
                                "tabs": {
                                    "signHereTabs":[{
                                        "anchorString": "/sn_customer/", #Placeholder for customer signature
                                        "anchorUnits": "pixels",
                                        "anchorXOffset": "0",
                                        "anchorYOffset": "0"
                                    }],
                                    "dateSignedTabs":[{
                                        "anchorString": "/ds_customer/", #Placeholder for customer signed date
                                        "anchorUnits": "pixels",
                                        "anchorXOffset": "0",
                                        "anchorYOffset": "0"
                                    }]
                                }
                            }
                        ],
                    },
                    #Tells DocuSign to send the envelop immediately for signing when the contract is created
                    "status": "sent"
                }
                #Extracts the ContactsToBeCCd and UsersToBeCCd values from the request body dictionary
                ContactsToBeCCd = data.get("ContactsToBeCCd")
                UsersToBeCCd = data.get("UsersToBeCCd")
                #Extracts the entities to be CCd, accumulates them into a list in the format required by docusign to be added to the envelope definition
                CarbonCopies = []
                accumulator = 2
                if ContactsToBeCCd:
                    for contact in ContactsToBeCCd:
                        accumulator += 1
                        if contact['ClientType'] == "Client_Company":
                            name = contact['name']
                            email = contact['CompanyEmail']
                        else:
                            name = f"{contact['Fname']} {contact['Lname']}"
                            email = contact['ClientEmail']
                        Contact = {"email": email, "name": name, "recipientId": str(accumulator), "routingOrder": str(accumulator)}
                        CarbonCopies.append(Contact)
                if UsersToBeCCd:
                    for user in UsersToBeCCd:
                        accumulator += 1
                        User = {"email": user['email'], "name": f"{user['Fname']} {user['Lname']}", "recipientId": str(accumulator), "routingOrder": str(accumulator)}
                        CarbonCopies.append(User)
                #adds recipients to be ccd in the contract email sent by docusign
                if CarbonCopies:
                    envelope_definition["recipients"].update({"carbonCopies":CarbonCopies})

                #Headers required for DocuSign API request
                headers = {
                    "Authorization": f"Bearer {access_token}", #Access token of the DocuSign account holder thats connected to my web app
                    "Content-Type": "application/json"
                }
                #DocuSign url to create an envelope
                url = f"{base_url}/accounts/{account_id}/envelopes"
                #Sends the request with the envelope definition, headers and the url as requires by DocuSign's API to create and send the contract
                response = requests.post(url, headers=headers, json=envelope_definition)
                response.raise_for_status()
                envelope_id = response.json()["envelopeId"]
                #Request payload to ask DocuSign to generate an embedded signing URL for the DocuSign account holder that is connected to my web app
                recipient_view_request_1 = {
                    "userName": director_name,
                    "email": director_email,
                    "recipientId": "1",
                    "clientUserId": "12345",
                    "authenticationMethod": "none",
                    "returnUrl": f"http://localhost:5173/Client/contracts?ClientType={client_type}&id={client_id}" #Return URL telling DocuSign to redirect the user to here (where they created the contract) after sining is completed
                }
                #DocuSign url to request the signing url
                view_url_1 = f"{base_url}/accounts/{account_id}/envelopes/{envelope_id}/views/recipient"
                #Sends the request for the embedded signing url
                view_response_1 = requests.post(view_url_1, headers=headers, json=recipient_view_request_1)
                view_response_1.raise_for_status() #Raises errors for debugging
                #Extracts the signing url from the response by docusign by indexing it
                signing_url_1 = view_response_1.json()["url"]
                #Modifies the created contract object by adding the signing url and envelope id to it
                contract.signingUrl = signing_url_1
                contract.envelope_id = envelope_id
                contract.save() #Saves the modified contract object to the database


            #Extracts the ProductsOrServices value from the request body dictionary
            ProductsOrServices = data.get("ProductsOrServices")
            #Loops over each item in the ProductsOrServices list
            for item in ProductsOrServices:
                #Creates a ContractItem object in the databasefor each product or service
                Contract_item = ContractItem.objects.create(Name=item['Name'], Price=item['Price'], Selected_VAT_Rate=item['Selected_VAT_Rate'], Total_price=item['Total_price'], VAT=item['VAT'], contract=contract, workplace=request.user.workplace)
                try:
                    #Extracts the deadline from the product or service item
                    Deadline = item['Deadline']
                    #Checks if the deadline vaue is not empty
                    if Deadline != '':
                        #Converts the deadline string into a date object
                        Deadline = datetime.strptime(Deadline, "%Y-%m-%d").date()
                        #If employees are assigned to work on the product or service creates relevant objects with the deadline in the databse for job distribution
                        if item['Employees']:
                            for user in item['Employees']:
                                job = Job.objects.create(corresponds_to_ContractItem=Contract_item, workplace=request.user.workplace, Deadline=Deadline)
                                user = CustomUser.objects.get(id=user['id'], workplace=request.user.workplace)
                                jobassigned = JobAssigned.objects.create(Job=job, User=user)
                        else:
                            pass
                    else:
                        #If employees are assigned to work on the product or service but there is no deadline provided creates relevant objects without the in the databse for job distribution
                        if item['Employees']:
                            for user in item['Employees']:
                                job = Job.objects.create(corresponds_to_ContractItem=Contract_item, workplace=request.user.workplace)
                                user = CustomUser.objects.get(id=user['id'], workplace=request.user.workplace)
                                jobassigned = JobAssigned.objects.create(Job=job, User=user)
                        else:
                            pass
                except:
                    pass
        else:
            client_individual = Client_Individual.objects.get(workplace=request.user.workplace, id=client_id)
            Signer = client_individual
            individual_fname = client_individual.Fname
            individual_Lname = client_individual.Lname
            individual_address = client_individual.address
            print(individual_address)
            Signer_First_Name = Signer.Fname
            print(Signer_First_Name)
            Signer_Last_Name = Signer.Lname
            print(Signer_Last_Name)
            #A dictionary with placeholders as keys and what they should be replaced with as their values
            replacements = {
                "{Client_First_Name}": individual_fname,
                "{Client_Last_Name}": individual_Lname,
                "{Client_Address}": individual_address,
                "{Client_First_Name}": Signer_First_Name,
                "{Client_Last_Name}": Signer_Last_Name,
                "{Directors_Signature}": "/sn_director/",
                "{Date_Signed_Director}": "/ds_director/",
                "{Client_Signature}": "/sn_customer/",
                "{Date_Signed_Client}": "/ds_customer/",
                "{Total_price}": Total['Price'],
                "{Total_Price_with_VAT}": Total['Price_with_VAT'],
            }
            print(replacements)
            #Gets the contract template from the database to be used as the input file in replace_placeholders function
            input_file = UploadedContractWordDocumentIndividual.objects.get(workplace=request.user.workplace)
            #Creates a contract object in the database
            contract = Contract.objects.create(workplace=request.user.workplace, client_individual=client_individual, Price=Total["Price"], VAT=Total["VAT"], total_amount=Total["Price_with_VAT"], People=People)
            #Builds different payments breakdowns strings and adds them to the replacements dictionary based on different invoice schedule conditions
            if number_of_invoices <= 1:
                replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be paid in 1 installment and will payable on {contract.created_at.strftime('%d-%m-%Y')}."})
            elif number_of_invoices > 1 and last_invoice_total_price:
                if invoice_duration == "for 1 year":
                    end_date = contract.created_at + relativedelta(years=1)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. Per instalment amount of the first {number_of_invoices - 1} instalments is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}. The last instalment amount is {last_invoice_total_price} plus VAT = {last_invoice_Total_with_VAT}."})
                elif invoice_duration == "for 1 week":
                    end_date = contract.created_at + timedelta(weeks=1)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. Per instalment amount of the first {number_of_invoices - 1} instalments is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}. The last instalment amount is {last_invoice_total_price} plus VAT = {last_invoice_Total_with_VAT}."})
                elif invoice_duration == "for 1 month":
                    end_date = contract.created_at + relativedelta(months=1)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. Per instalment amount of the first {number_of_invoices - 1} instalments is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}. The last instalment amount is {last_invoice_total_price} plus VAT = {last_invoice_Total_with_VAT}."})
                elif invoice_duration == "for 1 quarter":
                    end_date = contract.created_at + relativedelta(months=3)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. Per instalment amount of the first {number_of_invoices - 1} instalments is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}. The last instalment amount is {last_invoice_total_price} plus VAT = {last_invoice_Total_with_VAT}."})
                elif invoice_duration == "for 5 years":
                    end_date = contract.created_at + relativedelta(years=5)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. Per instalment amount of the first {number_of_invoices - 1} instalments is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}. The last instalment amount is {last_invoice_total_price} plus VAT = {last_invoice_Total_with_VAT}."})

            else:
                if invoice_duration == "for 1 year":
                    end_date = contract.created_at + relativedelta(years=1)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. {invoice_schedule} amount is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}."})
                elif invoice_duration == "for 1 week":
                    end_date = contract.created_at + timedelta(weeks=1)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. {invoice_schedule} amount is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}."})
                elif invoice_duration == "for 1 month":
                    end_date = contract.created_at + relativedelta(months=1)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. {invoice_schedule} amount is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}."})
                elif invoice_duration == "for 1 quarter":
                    end_date = contract.created_at + relativedelta(months=3)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. {invoice_schedule} amount is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}."})
                elif invoice_duration == "for 5 years":
                    end_date = contract.created_at + relativedelta(years=5)
                    end_date = end_date.strftime('%d-%m-%Y')
                    replacements.update({"{Payments_Breakdown}": f"The amount of {Total['Price']} plus VAT = {Total['Price_with_VAT']} will be divided into {number_of_invoices} {invoice_schedule} instalments and will be payable starting from {contract.created_at.strftime('%d-%m-%Y')} to {end_date}. {invoice_schedule} amount is {Total_Price_per_invoice} plus VAT = {Total_withVAT_per_invoice}."})
            #Calls the replace_placeholders function to replace placeholders in the contract template to create ready personalised contract and save it in the database (reference the replace_placeholders function)
            replace_placeholders(input_file, replacements, table_entries, contract, total_table_entries=Total)

            print(client_individual.email)

            try:
                DocuSign_subject = EmailSubjectforDocuSign_individual.objects.get(workplace = request.user.workplace)
                #Gets the ContentType object for the client individual
                content_type = ContentType.objects.get_for_model(client_individual)
                #Gets all the Represents objects from the database where the client individual is the object that the contacts are assigned to and where those contacts have a selected flag (True) because it means they are chosen for correspondence with the client by the user.
                represents_list = Represents.objects.filter(from_content_type=content_type, from_object_id = client_individual.id, workplace=request.user.workplace, selected=True)
                #Loops over the Represents objects list
                for rep in represents_list:
                    #Checks what client types the contacts are and retrives relevant information needed depending on what client types they are
                    if isinstance(rep.to_entity, Client_Individual):
                        Related_Contact_First_Name = rep.to_entity.Fname
                        Related_Contact_Last_Name = rep.to_entity.Lname
                    elif isinstance(rep.to_entity, Client_Company):
                        Related_Contact_Company_Name = rep.to_entity.name
                #replacements dictionary used to replace placeholders with the values as shown below
                replacements = {
                    "{Client_First_Name}": individual_fname,
                    "{Client_Last_Name}": individual_Lname,
                    "{Related_Contact_First_Name}": Related_Contact_First_Name,
                    "{Related_Contact_Last_Name}": Related_Contact_Last_Name,
                    "{Related_Contact_Company_Name}": Related_Contact_Company_Name,
                }
                #replaces placeholders in the email template from the database by calling the ReplacePlaceholdersCompany function and using the dictionary above (refer to the function)
                email_subject = ReplacePlaceholdersEmailCompany(DocuSign_subject.subject)
            except:
                email_subject = "Contract for Signature"
            #Opens the created contract file in read binary mode
            with open(contract.file.path, 'rb') as file:
                #Reads the file content into memory as bytes
                doc_bytes = file.read()
                #Encodes the file into based64 format and converts it to a utf-8 string because this is the required format for DocuSign API
                doc_base64 = base64.b64encode(doc_bytes).decode('utf-8')
                #Gets the DocuSign account detials from the database for the workplace the request is sent from
                DocuSignAccount = DocuSignToken.objects.get(workplace=request.user.workplace)
                access_token = DocuSignAccount.access_token
                account_id = DocuSignAccount.account_id
                director_email = DocuSignAccount.user_email
                director_name = DocuSignAccount.user_full_name
                base_url = "https://demo.docusign.net/restapi/v2.1" #API URL for DocuSign 
                #Builds the envelope definition to send to DocuSign api
                #The details are taken from DocuSign REST API documentation
                envelope_definition = {
                    "emailSubject": email_subject, #Subject of the email that the contract is sent in by DocuSign
                    #Documents to be included in the envelope (the contract in our case)
                    "documents": [{
                        "documentBase64": doc_base64, #The contract encoded in base64
                        "name": f"Contract_{contract.id}.docx", #Name of the contract
                        "fileExtension": "docx", #The file type of the project
                        "documentId": f"{contract.id}" #Unique id of the document to be sent (the contract)
                    }],
                    #People who will recieve the contract
                    "recipients": {
                        #People who will recieve the contract and will need to sign it
                        "signers": [
                            {
                                #First recipient: the account holder of the DocuSign account connected to my web app
                                "email": director_email,
                                "name": director_name,
                                "recipientId": "1", #Recipient id, has to be unique in this specific envelope only therefore I can just leave it at this
                                "routingOrder": "1", #The order for who gets the contract first, in this case it is the account holder of the DocuSign account connected to InfinitumAdmin
                                "clientUserId": "12345", #DocuSign requirement for embedded signing
                                #Signature and date placeholders for this recipient (the account holder). This is telling DocuSign that there are these placeholders in the document and telling it to search for them and place the fields over them
                                #Here I set all coordinated to 0 since I have this placeholder in the contract I am sending to the fields will be placed directly over it
                                "tabs": {
                                    "signHereTabs":[{
                                        "anchorString": "/sn_director/", #Placeholder for the DocuSign account holder signature
                                        "anchorUnits": "pixels",
                                        "anchorXOffset": "0",
                                        "anchorYOffset": "0"
                                    }],
                                    "dateSignedTabs":[{
                                        "anchorString": "/ds_director/", #Placeholder for the DocuSign account holder signed date
                                        "anchorUnits": "pixels",
                                        "anchorXOffset": "0",
                                        "anchorYOffset": "0"
                                    }]
                                }
                            },
                            {
                                #Second recipient:the client to whom the contract will be sent after the account holder signs it
                                "email": str(client_individual.email), 
                                "name": f"{Signer_First_Name} {Signer_Last_Name}",
                                "recipientId": "2", #Recipient id, has to be unique in this specific envelope only therefore I can just leave it at this
                                "routingOrder": "2", #The order for who gets the contract second, in this case it is the client who the contract is intended for by InfinitumAdmin user
                                #Signature and date placeholders for this recipient (the account holder). This is telling DocuSign that there are these placeholders in the document and telling it to search for them and place the fields over them
                                #Here I set all coordinated to 0 since I have this placeholder in the contract I am sending to the fields will be placed directly over it
                                "tabs": {
                                    "signHereTabs":[{
                                        "anchorString": "/sn_customer/", #Placeholder for customer signature
                                        "anchorUnits": "pixels",
                                        "anchorXOffset": "0",
                                        "anchorYOffset": "0"
                                    }],
                                    "dateSignedTabs":[{
                                        "anchorString": "/ds_customer/", #Placeholder for customer signed date
                                        "anchorUnits": "pixels",
                                        "anchorXOffset": "0",
                                        "anchorYOffset": "0"
                                    }]
                                }
                            }
                        ],
                    },
                    #Tells DocuSign to send the envelop immediately for signing when the contract is created
                    "status": "sent"
                }
                #Extracts the ContactsToBeCCd and UsersToBeCCd values from the request body dictionary
                ContactsToBeCCd = data.get("ContactsToBeCCd")
                UsersToBeCCd = data.get("UsersToBeCCd")
                #Extracts the entities to be CCd, accumulates them into a list in the format required by docusign to be added to the envelope definition
                CarbonCopies = []
                accumulator = 2
                if ContactsToBeCCd:
                    for contact in ContactsToBeCCd:
                        accumulator += 1
                        if contact['ClientType'] == "Client_company":
                            name = contact['CompanyName']
                            email = contact['CompanyEmail']
                        else:
                            name = f"{contact['Fname']} {contact['Lname']}"
                            email = contact['ClientEmail']
                        Contact = {"email": email, "name": name, "recipientId": str(accumulator), "routingOrder": str(accumulator)}
                        CarbonCopies.append(Contact)
                if UsersToBeCCd:
                    for user in UsersToBeCCd:
                        accumulator += 1
                        User = {"email": user['email'], "name": f"{user['Fname']} {user['Lname']}", "recipientId": str(accumulator), "routingOrder": str(accumulator)}
                        CarbonCopies.append(User)
                #adds recipients to be ccd in the contract email sent by docusign
                if CarbonCopies:
                    envelope_definition["recipients"].update({"cardbonCopies":CarbonCopies})

                #Headers required for DocuSign API request
                headers = {
                    "Authorization": f"Bearer {access_token}", #Access token of the DocuSign account holder thats connected to my web app
                    "Content-Type": "application/json"
                }
                #DocuSign url to create an envelope
                url = f"{base_url}/accounts/{account_id}/envelopes"
                #Sends the request with the envelope definition, headers and the url as requires by DocuSign's API to create and send the contract
                response = requests.post(url, headers=headers, json=envelope_definition)
                response.raise_for_status()
                envelope_id = response.json()["envelopeId"]

                recipient_view_request_1 = {
                    "userName": director_name,
                    "email": director_email,
                    "recipientId": "1",
                    "clientUserId": "12345",
                    "authenticationMethod": "none",
                    "returnUrl": f"http://localhost:5173/Client/contracts?ClientType={client_type}&id={client_id}" #Return URL telling DocuSign to redirect the user to here (where they created the contract) after sining is completed
                }

                view_url_1 = f"{base_url}/accounts/{account_id}/envelopes/{envelope_id}/views/recipient"
                view_response_1 = requests.post(view_url_1, headers=headers, json=recipient_view_request_1)
                view_response_1.raise_for_status() #Raises errors for debugging
                #Extracts the signing url from the response by docusign by indexing it
                signing_url_1 = view_response_1.json()["url"]
                #Modifies the created contract object by adding the signing url and envelope id to it
                contract.signingUrl = signing_url_1
                contract.envelope_id = envelope_id
                contract.save() #Saves the modified contract object to the database


            #Extracts the ProductsOrServices value from the request body dictionary
            ProductsOrServices = data.get("ProductsOrServices")
            #Loops over each item in the ProductsOrServices list
            for item in ProductsOrServices:
                #Creates a ContractItem object in the databasefor each product or service
                Contract_item = ContractItem.objects.create(Name=item['Name'], Price=item['Price'], Selected_VAT_Rate=item['Selected_VAT_Rate'], Total_price=item['Total_price'], VAT=item['VAT'], contract=contract, workplace=request.user.workplace)
                try:
                    #Extracts the deadline from the product or service item
                    Deadline = item['Deadline']
                    #Checks if the deadline vaue is not empty
                    if Deadline != '':
                        #Converts the deadline string into a date object
                        Deadline = datetime.strptime(Deadline, "%Y-%m-%d").date()
                        if item['Employees']:
                            #If employees are assigned to work on the product or service creates relevant objects with the deadline in the databse for job distribution
                            for user in item['Employees']:
                                job = Job.objects.create(corresponds_to_ContractItem=Contract_item, workplace=request.user.workplace, Deadline=Deadline)
                                user = CustomUser.objects.get(id=user['id'], workplace=request.user.workplace)
                                jobassigned = JobAssigned.objects.create(Job=job, User=user)
                        else:
                            pass
                    else:
                        if item['Employees']:
                            #If employees are assigned to work on the product or service but there is no deadline provided creates relevant objects without the in the databse for job distribution
                            for user in item['Employees']:
                                job = Job.objects.create(corresponds_to_ContractItem=Contract_item, workplace=request.user.workplace)
                                user = CustomUser.objects.get(id=user['id'], workplace=request.user.workplace)
                                jobassigned = JobAssigned.objects.create(Job=job, User=user)
                        else:
                            pass
                except:
                    pass
        #Returns signing url and contract json object to the front end as a response
        return JsonResponse({"signing_url_director": signing_url_1, "contract": contract.as_dict()})
    
    if request.method == 'PUT':
        #Retrieves the request data and turns it into python readable dictionary
        data = json.loads(request.body)
        #Extracts the contract, id and Signed values from the request body dictionary
        contract_data = data.get('contract')
        contract_id = contract_data.get('id')
        Signed = contract_data.get('Signed')
        #Gets the contract object from the database
        contract = Contract.objects.get(id = contract_id, workplace = request.user.workplace)
        #Assigns the Signed value to the contract.Signed attribute
        contract.Signed = Signed
        contract.save() #Saves the modified contract object in the database
        return JsonResponse(contract.as_dict())
    
    if request.method == 'DELETE':
        #Retrieves the request data and turns it into python readable dictionary
        data = json.loads(request.body)
        #Extracts the contract_id value from the request body dictionary
        contract_id = data.get('contract_id')
        #Gets the contract object from the database
        contract = Contract.objects.get(id = contract_id, workplace = request.user.workplace)
        #Sets the number_of_invoices summary of the Invoice_Summary related to this contract to 0 because by the code logic in the cron.py file all invoice automations related to this contract will be skipped
        try:
            invoice_summary = Invoice_Summary.objects.get(contract = contract, workplace = request.user.workplace)
            invoice_summary.number_of_invoices = 0
            invoice_summary.save()
        except Invoice_Summary.DoesNotExist:
            pass
        contract.delete() #Deletes the contract object
        return JsonResponse({'DELETE': 'Confirm'})

#This function replaces placeholders in a Word document template with values which are passed to it
def replace_placeholders_invoice(input_file, replacements, table_entries, invoice, total_table_entries):
    #Loads the word file which is going to be edited
    doc = Document(input_file.file.path)
    placeholder_regex = r"\{.*?\}" #Regular expression to match any text inside the {} placeholders

    #This function replaces placeholders in a paragraph
    def replace_in_paragraph(paragraph, replace_dict=None):
        #This assigns the passed replacements dictionary to be used as replacements if it is passed, otherwise it uses replacements passed in the parent function.
        current_replacements = replace_dict if replace_dict else replacements
        full_text = paragraph.text
        #This replaces the placeholders in the paragraph using the replacements dictionary passed to this function is any placeholders are matched by the regular expression
        if re.search(placeholder_regex, full_text):
            #This loops over all keys (placeholders) and their values to replace them in the paragraph
            for placeholder, value in current_replacements.items():
                full_text = full_text.replace(str(placeholder), str(value))
            #This clears all existing text runs in the paragraph
            for run in paragraph.runs:
                run.text = ""
            #If the passed paragraph has runs, this replaces the original text with placeholders with the text which has palceholders replaced 
            if paragraph.runs:
                paragraph.runs[0].text = full_text
                run=paragraph.runs[0]
            else:
                #If no runs exist this adds a new run with the replaces text
                paragraph.add_run(full_text)
    #Here I loop over all paragraphs in the document and replace placeholders in them using the function above 
    for para in doc.paragraphs:
        replace_in_paragraph(para)
    #Here I loop over all tables in the document
    for table in doc.tables:
        #In every table I iterate over each row and each cell in each row and each paragraph in each cell
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    #If any placeholder from the passed dictionaty is in the paragraph I replace it in the paragraph using the the replacements dictionary
                    if any(str(placeholder) in para.text for placeholder in replacements.keys()):
                        replace_in_paragraph(para, replace_dict=replacements)
    #Here I loop over through all tables to insert relevant values replacement placeholder fiels in them
    for table in doc.tables:
        #In every table I iterate over each row and each cell in each row and each paragraph in each cell
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    #If any placeholder from the passed dictionaty is in the paragraph I replace it in the paragraph using the the replacements dictionary
                    if any(str(placeholder) in para.text for placeholder in replacements.keys()):
                        replace_in_paragraph(para, replace_dict=replacements)
    #I iterate over each dictionary key and value pair in the table entries dictionary passed to this function
    for entry in table_entries:
        #I replace the original keys with placeholders I want to be there I do this because I pass the original keys in table entries like that in the api_Invoice function when I call this function
        replace_key(entry, 'Name', '{Service_or_Product_Name}')
        replace_key(entry, 'Total_price', '{Price_with_VAT}')
        replace_key(entry, 'Price', '{Price}')
        replace_key(entry, 'VAT', '{VAT}')
    #Here I loop over through all tables to insert relevant values replacement placeholder fiels in them
    for table in doc.tables:
        #Here I initialise a template row variable to copy a template row into it 
        template_row = None
        #Here I iterate over all rows in the table after the first row (I assume it is the header)
        for row in table.rows[1:]:
            found = False
            #I iterate over each cell in each row and each paragraph in each cell
            for cell in row.cells:
                for para in cell.paragraphs:
                    for entry in table_entries:
                        if isinstance(entry, dict):
                            #If any of the placeholders from table entries exist in the paragraph
                            if any(str(placeholder) in para.text for placeholder in entry.keys()):
                                #Then the template row is found
                                template_row = row
                                found = True
                                break
                    if found:
                        break
                if found:
                    break
            if found:
                    break
        #If template row is found
        if template_row:
            #Loops over all table entries
            for entry in table_entries:
                #Deep copies the template row to create a new row for each entry in the table entries and at the same time avoid modifying the found row
                new_row = deepcopy(template_row)
                #Loops over each cell and each paragraph in each cell of the new row
                for cell in new_row.cells:
                    for para in cell.paragraphs:
                        #replaces each paragraph with the replacement dictionary
                        replace_in_paragraph(para, replace_dict=entry)
                #Inserts a new row
                idx = table._tbl.index(template_row._tr)
                table._tbl.insert(idx+1, new_row._tr)
            #removes the original found row in the table
            table._tbl.remove(template_row._tr)
    #Gets the last row of the table (I assume this is where placeholders for total amount are)
    last_row = table.rows[-1]
    #I replace the original keys with placeholders I want to be there I do this because I pass the original keys in table entries like that in the api_Invoice function when I call this function
    replace_key(total_table_entries, 'Price', '{Total_price}')
    replace_key(total_table_entries, 'VAT', '{Total_VAT}')
    replace_key(total_table_entries, 'Price_with_VAT', '{Total_Price_with_VAT}')
    #Replaces placeholders in the last row with the actual totals values
    for cell in last_row.cells:
        for para in cell.paragraphs:
            replace_in_paragraph(para, replace_dict=total_table_entries)

    #Creates a temporary file with a .docx suffix to store the word document temporarily
    tmp_file =  tempfile.NamedTemporaryFile(suffix=".docx", delete=True)
    tmp_file_name = tmp_file.name #Saves the temporary file path
    tmp_file.close() #Closes the file
    doc.save(tmp_file_name) #Saves the current word document into the temporary file created
    

    output_dir = os.path.dirname(tmp_file_name) #Gets the directory where the temporary file is stored
    #Defines the path to LibreOffice soffice.exe which is used for converting docx files to pdf files
    libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
    #Runs LibreOffice in headless mode to convert the file to pdf
    subprocess.run([
        libreoffice_path,
        '--headless', #Indicates to the subprocess to use LibraOffice in headless mode
        '--convert-to', 'pdf', #Tells subprocess to convert the file to pdf
        '--outdir', output_dir, #Tells subprocess where to put the pdf file
        tmp_file_name #The temporary docx file to be converted
    ], check=True) #Raises an error if LibraOffice error occurs
    #Changes the .docx extention of the file name to .pdf
    pdf_filename = tmp_file_name.replace('.docx', '.pdf')
    #If the PSD file exists, opens it in read binary mode and saves it into the invoice object file path
    if os.path.exists(pdf_filename):
        with open(pdf_filename, 'rb') as f:
            invoice.invoice_file.save(f"invoice_{invoice.id}.pdf", File(f), save=True)
    #Removes the temporary word file
    if os.path.exists(tmp_file_name):
        os.remove(tmp_file_name)
    #Removes the pdf file
    if os.path.exists(pdf_filename):
        os.remove(pdf_filename)
            
#This function is used to handle client requests to get invoices information, create an invoice and delete an invoice and all associated automations with it.
@csrf_exempt
@login_required
def api_Invoice(request):
    if request.method == 'GET':
        if request.headers.get('Url-Header'):
            #Getting the header
            url = request.headers.get('Url-Header')
            parsed_url = urlparse(url) #splits the url into its components (the path, quey, fragment etc..) that I can access
            query_params = parse_qs(parsed_url.query) #converts the query string of the url into a dictionary with keys being single strings such as 'ClientType' and values being lists of values
            client_id = query_params.get('id')[0] #extracting the client id from the url by indexing the value of the 'id' key as [0] because the values are stored in a list
            client_type = query_params.get('ClientType')[0] #extracting the client type from the url by indexing the value of the 'ClientType' key as [0] because the values are stored in a list
            #Returns a Json response to the front end based on the url-header
            if client_type == 'Client_company':
                return JsonResponse({'InvoicesToInvoiceItems': [invoice.as_dict() for invoice in InvoicesToInvoiceItems.objects.filter(workplace=request.user.workplace, invoice__client_company__id=client_id)]})
            else:
                return JsonResponse({'InvoicesToInvoiceItems': [invoice.as_dict() for invoice in InvoicesToInvoiceItems.objects.filter(workplace=request.user.workplace, invoice__client_individual__id=client_id)]})
        else:
            return JsonResponse({
                'InvoicesToInvoiceItems': [invoice.as_dict() for invoice in InvoicesToInvoiceItems.objects.filter(workplace = request.user.workplace)]
            })
    
    elif request.method == 'POST':
        #Getting the header
        url = request.headers.get('Url-Header')
        parsed_url = urlparse(url) #splits the url into its components (the path, quey, fragment etc..) that I can access
        query_params = parse_qs(parsed_url.query) #converts the query string of the url into a dictionary with keys being single strings such as 'ClientType' and values being lists of values
        client_id = query_params.get('id')[0] #extracting the client id from the url by indexing the value of the 'id' key as [0] because the values are stored in a list
        client_type = query_params.get('ClientType')[0] #extracting the client type from the url by indexing the value of the 'ClientType' key as [0] because the values are stored in a list

        data = json.loads(request.body) #Retrieves the request body and turns it into python readable dictionary
        invoice_schedule_summary = data.get("Invoice_Schedule") #Extracts the Invoice_Schedule value from the request body dictionary
        #Extracts various values from the InvoiceSchedule dictionary
        number_of_invoices = invoice_schedule_summary['number_of_invoices']
        schedule = invoice_schedule_summary['invoice_schedule']
        duration = invoice_schedule_summary['duration'] 
        Total_price = invoice_schedule_summary['totalPrice']
        Total_VAT = invoice_schedule_summary['totalVAT']
        Total_with_VAT = invoice_schedule_summary['totalWithVAT']

        reminders = invoice_schedule_summary['reminders']
        #Extracts count and per values from the reminders dictionary 
        reminders_count = reminders['count']
        reminders_per = reminders['per']
        #Checks if the invoice schedule relates to a contract
        if data.get("Contract"):
            #Gets the contract id
            contract_id = data.get("Contract")['id']
            print("Success")
            #Gets the contract object from the database
            contract = Contract.objects.get(id=contract_id, workplace=request.user.workplace)
            if client_type == 'Client_company':
                Client_company = Client_Company.objects.get(id=client_id)
                #Creates the invoice summary object in the database
                try:
                    last_invoice = invoice_schedule_summary['lastInvoice']
                    last_invoice_Total_price = last_invoice['totalPrice']
                    last_invoice_Total_VAT = last_invoice['totalVAT']
                    last_invoice_Total_with_VAT = last_invoice['totalWithVAT']
                    invoice_summary = Invoice_Summary.objects.create(number_of_invoices = number_of_invoices, schedule=schedule, duration=duration, reminders_count=reminders_count, reminders_per=reminders_per, contract=contract, Total_price=Total_price, Total_VAT=Total_VAT, Total_with_VAT=Total_with_VAT, last_invoice_Total_price=last_invoice_Total_price, last_invoice_Total_VAT=last_invoice_Total_VAT, last_invoice_Total_with_VAT=last_invoice_Total_with_VAT, workplace=request.user.workplace, billed_to_company = Client_company)
                except:
                    invoice_summary = Invoice_Summary.objects.create(number_of_invoices = number_of_invoices, schedule=schedule, duration=duration, reminders_count=reminders_count, reminders_per=reminders_per, contract=contract, Total_price=Total_price, Total_VAT=Total_VAT, Total_with_VAT=Total_with_VAT, workplace=request.user.workplace, billed_to_company = Client_company)
                #Creates a unique identifier for the invoice (I thought that concatenating the company name with current date time down to seconds would guarantee uniqueness of invoice ID as well as serve as a good reference) (this is a personal choice)
                Company_name = (str(Client_company.name)).replace(' ', '')
                Current_date_time = timezone.now().strftime('%d%m%Y%H%M%S')
                identifier = (Company_name + Current_date_time).replace(" ", "")
                #Creates the invoice object in the database
                invoice = Invoice.objects.create(InvoiceSummary=invoice_summary, workplace = request.user.workplace, client_company=Client_company, identifier=identifier, contract=contract)
                print(invoice)
            else:
                Client_individual = Client_Individual.objects.get(id=client_id)
                #Creates the invoice summary object in the database
                try:
                    last_invoice = invoice_schedule_summary['lastInvoice']
                    last_invoice_Total_price = last_invoice['totalPrice']
                    last_invoice_Total_VAT = last_invoice['totalVAT']
                    last_invoice_Total_with_VAT = last_invoice['totalWithVAT']
                    invoice_summary = Invoice_Summary.objects.create(number_of_invoices = number_of_invoices, schedule=schedule, duration=duration, reminders_count=reminders_count, reminders_per=reminders_per, contract=contract, Total_price=Total_price, Total_VAT=Total_VAT, Total_with_VAT=Total_with_VAT, last_invoice_Total_price=last_invoice_Total_price, last_invoice_Total_VAT=last_invoice_Total_VAT, last_invoice_Total_with_VAT=last_invoice_Total_with_VAT, workplace=request.user.workplace, billed_to_individual = Client_individual)
                except:
                    invoice_summary = Invoice_Summary.objects.create(number_of_invoices = number_of_invoices, schedule=schedule, duration=duration, reminders_count=reminders_count, reminders_per=reminders_per, contract=contract, Total_price=Total_price, Total_VAT=Total_VAT, Total_with_VAT=Total_with_VAT, workplace=request.user.workplace, billed_to_individual = Client_individual)
                #Creates a unique identifier for the invoice (I thought that concatenating the company name with current date time down to seconds would guarantee uniqueness of invoice ID as well as serve as a good reference) (this is a personal choice)
                Individual_Fname = (str(Client_individual.Fname))
                Individual_Lname = (str(Client_individual.Lname))
                Current_date_time = timezone.now().strftime('%d%m%Y%H%M%S')
                identifier = (Individual_Fname+Individual_Lname+Current_date_time).replace(" ", "")
                #Creates the invoice object in the database
                invoice = Invoice.objects.create(InvoiceSummary=invoice_summary, workplace = request.user.workplace, client_individual=Client_individual, identifier=identifier, contract=contract)
                print(invoice)
            #Extracts information about entities who the user configured to send emails with invoices to
            ContactsToBeTod = data.get("ContactsToBeTod")
            ContactsToBeCCd = data.get("ContactsToBeCCd")
            ContactsToBeBccd = data.get("ContactsToBeBccd")
            UsersToBeTod = data.get("UsersToBeTod")
            UsersToBeCCd = data.get("UsersToBeCCd")
            UsersToBeBccd = data.get("UsersToBeBccd")
            MainClientToBeTod = data.get("MainClientToBeTod")
            MainClientToBeCCd = data.get("MainClientToBeCCd")
            MainClientToBeBccd = data.get("MainClientToBeBccd")
            #Builds lists of entities to be included as different types of email recipients
            to = []
            cc = []
            bcc = []
            if client_type == 'Client_company':
                Client_company = Client_Company.objects.get(id=client_id)
                if MainClientToBeTod == True:
                    invoice_summary.MainClientToBeTod = True
                    invoice_summary.save()
                    to.append(Client_company.email)
                elif MainClientToBeCCd == True:
                    invoice_summary.MainClientToBeCCd = True
                    invoice_summary.save()
                    cc.append(Client_company.email)
                elif MainClientToBeBccd == True:
                    invoice_summary.MainClientToBeBccd = True
                    invoice_summary.save()
                    bcc.append(Client_company.email)
                else:
                    pass
            else:
                Client_individual = Client_Individual.objects.get(id=client_id)
                if MainClientToBeTod == True:
                    invoice_summary.MainClientToBeTod = True
                    invoice_summary.save()
                    to.append(Client_individual.email)
                elif MainClientToBeCCd == True:
                    invoice_summary.MainClientToBeCCd = True
                    invoice_summary.save()
                    cc.append(Client_individual.email)
                elif MainClientToBeBccd == True:
                    invoice_summary.MainClientToBeBccd = True
                    invoice_summary.save()
                    bcc.append(Client_individual.email)
                else:
                    pass
                
            if UsersToBeTod:
                for user in UsersToBeTod:
                    user = CustomUser.objects.get(id=user['id'], workplace=request.user.workplace)
                    User = UsersToBeTodThrough.objects.create(invoice_summary = invoice_summary, user = user, workplace = request.user.workplace)
                    to.append(User.user.email)
            if UsersToBeCCd:
                for user in UsersToBeCCd:
                    user = CustomUser.objects.get(id=user['id'], workplace=request.user.workplace)
                    User = UsersToBeCCdThrough.objects.create(invoice_summary = invoice_summary, user = user, workplace = request.user.workplace)
                    cc.append(User.user.email)
            if UsersToBeBccd:
                for user in UsersToBeBccd:
                    user = CustomUser.objects.get(id=user['id'], workplace=request.user.workplace)
                    User = UsersToBeBccdThrough.objects.create(invoice_summary = invoice_summary, user = user, workplace = request.user.workplace)
                    bcc.append(User.user.email)
            if ContactsToBeTod:
                for contact in ContactsToBeTod:
                    if contact['ClientType'] == 'Client_Company':
                        client_company = Client_Company.objects.get(id=contact['id'], workplace=request.user.workplace)
                        Contact = ContactsCompaniesToBeTodThrough.objects.create(invoice_summary=invoice_summary, client_company = client_company, workplace=request.user.workplace)
                        to.append(Contact.client_company.email)
                    else:
                        client_individual = Client_Individual.objects.get(id=contact['id'], workplace = request.user.workplace)
                        Contact = ContactsIndividualsToBeTodThrough.objects.create(invoice_summary=invoice_summary, client_individual=client_individual, workplace=request.user.workplace)
                        to.append(Contact.client_individual.email)
            if ContactsToBeCCd:
                for contact in ContactsToBeCCd:
                    if contact['ClientType'] == 'Client_Company':
                        client_company = Client_Company.objects.get(id=contact['id'], workplace=request.user.workplace)
                        Contact = ContactsCompaniesToBeCCdThrough.objects.create(invoice_summary=invoice_summary, client_company = client_company, workplace=request.user.workplace)
                        cc.append(Contact.client_company.email)
                    else:
                        client_individual = Client_Individual.objects.get(id=contact['id'], workplace = request.user.workplace)
                        Contact = ContactsIndividualsToBeCCdThrough.objects.create(invoice_summary=invoice_summary, client_individual=client_individual, workplace=request.user.workplace)
                        cc.append(Contact.client_individual.email)
            if ContactsToBeBccd:
                for contact in ContactsToBeBccd:
                    if contact['ClientType'] == 'Client_Company':
                        client_company = Client_Company.objects.get(id=contact['id'], workplace=request.user.workplace)
                        Contact = ContactsCompaniesToBeBccdThrough.objects.create(invoice_summary=invoice_summary, client_company = client_company, workplace=request.user.workplace)
                        bcc.append(Contact.client_company.email)
                    else:
                        client_individual = Client_Individual.objects.get(id=contact['id'], workplace = request.user.workplace)
                        Contact = ContactsIndividualsToBeBccdThrough.objects.create(invoice_summary=invoice_summary, client_individual=client_individual, workplace=request.user.workplace)
                        bcc.append(Contact.client_individual.email)


            #Adds entries of invoice items in the database and creates a list of them
            invoice_items = []
            products = invoice_schedule_summary['products']
            for product in products:
                ProductName = product['Name']
                Price = product['PricePerInvoice']
                VAT = product['VATPerInvoice']
                Total_PRICE = product['TotalPerInvoice']
                invoice_item = InvoiceItem.objects.create(Name = ProductName, Price = Price, VAT = VAT, Total_price = Total_PRICE, invoice_summary = invoice_summary)
                invoice_items.append(invoice_item)
            try:
                last_invoice_products = last_invoice['products']
                for product in last_invoice_products:
                    ProductName = product['Name']
                    Price = product['PricePerInvoice']
                    VAT = product['VATPerInvoice']
                    Total_PRICE = product['TotalPerInvoice']
                    #The last=True flag in the InvoiceItem object specifies that these entries should be used in the last invoice as it could be different from the rest
                    invoice_item = InvoiceItem.objects.create(Name = ProductName, Price = Price, VAT = VAT, Total_price = Total_PRICE, invoice_summary = invoice_summary, last=True)
                    invoice_items.append(invoice_item)
            except:
                pass
            #Creates InvoicesToInvoiceItems object in the database to link the invoice with invoice items in the database
            InvoicesTo_invoice_items = []
            for invoice_item in invoice_items:
                invoices_to_invoice_items = InvoicesToInvoiceItems.objects.create(invoice_item = invoice_item, invoice=invoice, workplace=request.user.workplace)
                InvoicesTo_invoice_items.append(invoices_to_invoice_items.as_dict())
            print("Success")
            if client_type == 'Client_company':
                print("Success")
                company_name = Client_company.name
                company_address = Client_company.address
                input_file = UploadedInvoiceWordDocumentCompany.objects.get(workplace=request.user.workplace)
                Date = datetime.today().strftime('%d-%m-%Y')
                #The due date of the invoice is set by me on the same day the invoice is created as this is a prototype application
                Due_date = datetime.today().strftime('%d-%m-%Y')
                Invoice_Number = invoice.identifier
                #Builds a table entries list to be used in the replace_placeholders_invoice function to populate the invoice with relevant data
                table_entries = []

                for item in invoice_items:
                    if item.last == False:
                        entry = {'Name': item.Name, 'Price': item.Price, 'VAT': item.VAT, 'Total_price': item.Total_price}
                        table_entries.append(entry)
                print(table_entries)
                
                Total_price = invoice_summary.Total_price
                Total_VAT = invoice_summary.Total_VAT
                Total_with_VAT = invoice_summary.Total_with_VAT
                #Table total table entries list to be used to replace the total values in the invoice
                total_table_entries = {'Price': Total_price, 'VAT': Total_VAT, 'Price_with_VAT': Total_with_VAT}
                #Replacements dictionary to be used to replace placeholders with relevant information in the invoice using the replace_placeholders_invoice function
                replacements = {
                "{Company_Name}": company_name,
                "{Company_Address}": company_address,
                "{Date}": Date,
                "{Invoice_Number}": Invoice_Number,
                "{Due_date}": Due_date
                }
                #Creates the invoice using the replace_placeholders_invoice function
                replace_placeholders_invoice(input_file, replacements, table_entries, invoice, total_table_entries)
                #Creates a dictionary used for replacements for placeholders in the email that the invoice will be attached to
                if to or cc or bcc:
                        Email_template = EmailInvoiceFirstWithContract_company.objects.get(workplace=request.user.workplace)
                        content_type = ContentType.objects.get_for_model(Client_company)
                        Related_Contact_First_Name = None
                        Related_Contact_Last_Name = None
                        Related_Contact_Company_Name = None
                        represents_list = Represents.objects.filter(from_content_type=content_type, from_object_id = Client_company.id, workplace=request.user.workplace, selected=True)
                        for rep in represents_list:
                            if isinstance(rep.to_entity, Client_Individual):
                                Related_Contact_First_Name = rep.to_entity.Fname
                                Related_Contact_Last_Name = rep.to_entity.Lname
                            elif isinstance(rep.to_entity, Client_Company):
                                Related_Contact_Company_Name = rep.to_entity.name
                        Signer = Client_company.Signer
                        Signer_First_Name = Signer.Fname
                        Signer_Last_Name = Signer.Lname
                        #The value of {Period} placeholder is chosen by me
                        if invoice_summary.schedule == "One Invoice":
                            current_date = time.strftime("%d-%m-%Y")
                            invoice_schedule = current_date
                        elif invoice_summary.schedule == "Daily":
                            current_date = time.strftime("%d-%m-%Y")
                            invoice_schedule = current_date
                        elif invoice_summary.schedule == "Weekly":
                            today = time.localtime()
                            week_number = int(time.strftime("%w", today))
                            weekday = int(time.strftime("%u", today))
                            seconds_in_day = 86400
                            today_timestamp = time.mktime(today)
                            days_since_monday = weekday - 1
                            monday_timestamp = today_timestamp - (days_since_monday * seconds_in_day)
                            monday = time.localtime(monday_timestamp)
                            today_str = time.strftime("%d-%m-%Y", today)
                            monday_str = time.strftime("%d-%m-%Y", monday)
                            invoice_schedule = f"week staring on {monday_str}"
                        elif invoice_summary.schedule == "Monthly":
                            current_month_date = time.strftime("%B-%Y")
                            invoice_schedule = current_month_date
                        elif invoice_summary.schedule == "Quarterly":
                            month = int(time.strftime("%m"))
                            quarter = (month - 1) // 3 + 1
                            current_year_date = time.strftime("%Y")
                            invoice_schedule = f"q{quarter} of {current_year_date}"
                        elif invoice_summary.schedule == "Yearly":
                            current_year_date = time.strftime("%Y")
                            invoice_schedule = current_year_date
                        replacements = {
                            "{Company_Name}": company_name,
                            "{Signer_First_Name}": Signer_First_Name,
                            "{Signer_Last_Name}": Signer_Last_Name,
                            "{Period}": invoice_schedule,
                            "{Related_Contact_First_Name}": Related_Contact_First_Name,
                            "{Related_Contact_Last_Name}": Related_Contact_Last_Name,
                            "{Related_Contact_Company_Name}": Related_Contact_Company_Name,
                        }
                        #Replaces the placeholders in the email template subject and body
                        email_subject = ReplacePlaceholdersEmailCompany(Email_template.subject, replacements)
                        email_body = ReplacePlaceholdersEmailCompany(Email_template.body, replacements)
                        Subject = email_subject
                        Body = email_body
                        #Builds an email message based on which recipients are included in the email
                        if to and cc and bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to, cc=cc, bcc=bcc)
                        elif to and cc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to, cc=cc)
                        elif to and bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to, bcc=bcc)
                        elif cc and bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, cc = cc, bcc=bcc)
                        elif to:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to)
                        elif cc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, cc = cc)
                        elif bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, bcc = bcc)
                        #Attaches the invoice pdf to the email
                        email.attach_file(invoice.invoice_file.path)
                        #Sends the email
                        email.send()
                #Reduces the number of invoice in invoice summary by one because one invoice was already sent (used for cron jobs tracking)
                invoice_summary.number_of_invoices -= 1
                invoice_summary.save()

            else:
                print("Success")
                individual_name = f"{Individual_Fname} {Individual_Lname}"
                individual_address = Client_individual.address
                input_file = UploadedInvoiceWordDocumentIndividual.objects.get(workplace=request.user.workplace)
                Date = datetime.today().strftime('%d-%m-%Y')
                #The due date of the invoice is set by me on the same day the invoice is created as this is a prototype application
                Due_date = datetime.today().strftime('%d-%m-%Y')
                Invoice_Number = invoice.identifier
                #Builds a table entries list to be used in the replace_placeholders_invoice function to populate the invoice with relevant data
                table_entries = []
                for item in invoice_items:
                    if item.last == False:
                        entry = {'Name': item.Name, 'Price': item.Price, 'VAT': item.VAT, 'Total_price': item.Total_price}
                        table_entries.append(entry)
                print(table_entries)
                
                Total_price = invoice_summary.Total_price
                Total_VAT = invoice_summary.Total_VAT
                Total_with_VAT = invoice_summary.Total_with_VAT
                #Table total table entries list to be used to replace the total values in the invoice
                total_table_entries = {'Price': Total_price, 'VAT': Total_VAT, 'Price_with_VAT': Total_with_VAT}
                #Replacements dictionary to be used to replace placeholders with relevant information in the invoice using the replace_placeholders_invoice function
                replacements = {
                "{Client_First_Name}": Individual_Fname,
                "{Client_Last_Name}": Individual_Lname,
                "{Client_Address}": individual_address,
                "{Date}": Date,
                "{Invoice_Number}": Invoice_Number,
                "{Due_date}": Due_date
                }
                #Creates the invoice using the replace_placeholders_invoice function
                replace_placeholders_invoice(input_file, replacements, table_entries, invoice, total_table_entries)
                #Creates a dictionary used for replacements for placeholders in the email that the invoice will be attached to
                if to or cc or bcc:
                        Email_template = EmailInvoiceFirstWithContract_individual.objects.get(workplace=request.user.workplace)
                        content_type = ContentType.objects.get_for_model(Client_individual)
                        Related_Contact_First_Name = None
                        Related_Contact_Last_Name = None
                        Related_Contact_Company_Name = None
                        represents_list = Represents.objects.filter(from_content_type=content_type, from_object_id = Client_individual.id, workplace=request.user.workplace, selected=True)
                        for rep in represents_list:
                            if isinstance(rep.to_entity, Client_Individual):
                                Related_Contact_First_Name = rep.to_entity.Fname
                                Related_Contact_Last_Name = rep.to_entity.Lname
                            elif isinstance(rep.to_entity, Client_Company):
                                Related_Contact_Company_Name = rep.to_entity.name
                        Signer = Client_individual
                        Signer_First_Name = Signer.Fname
                        Signer_Last_Name = Signer.Lname
                        #The value of {Period} placeholder is chosen by me
                        if invoice_summary.schedule == "One Invoice":
                            current_date = time.strftime("%d-%m-%Y")
                            invoice_schedule = current_date
                        elif invoice_summary.schedule == "Daily":
                            current_date = time.strftime("%d-%m-%Y")
                            invoice_schedule = current_date
                        elif invoice_summary.schedule == "Weekly":
                            today = time.localtime()
                            week_number = int(time.strftime("%w", today))
                            weekday = int(time.strftime("%u", today))
                            seconds_in_day = 86400
                            today_timestamp = time.mktime(today)
                            days_since_monday = weekday - 1
                            monday_timestamp = today_timestamp - (days_since_monday * seconds_in_day)
                            monday = time.localtime(monday_timestamp)
                            today_str = time.strftime("%d-%m-%Y", today)
                            monday_str = time.strftime("%d-%m-%Y", monday)
                            invoice_schedule = f"week staring on {monday_str}"
                        elif invoice_summary.schedule == "Monthly":
                            current_month_date = time.strftime("%B-%Y")
                            invoice_schedule = current_month_date
                        elif invoice_summary.schedule == "Quarterly":
                            month = int(time.strftime("%m"))
                            quarter = (month - 1) // 3 + 1
                            current_year_date = time.strftime("%Y")
                            invoice_schedule = f"q{quarter} of {current_year_date}"
                        elif invoice_summary.schedule == "Yearly":
                            current_year_date = time.strftime("%Y")
                            invoice_schedule = current_year_date
                        replacements = {
                            "{Client_First_Name}": Individual_Fname,
                            "{Client_Last_Name}": Individual_Lname,
                            "{Period}": invoice_schedule,
                            "{Related_Contact_First_Name}": Related_Contact_First_Name,
                            "{Related_Contact_Last_Name}": Related_Contact_Last_Name,
                            "{Related_Contact_Company_Name}": Related_Contact_Company_Name,
                        }
                        #Replaces the placeholders in the email template subject and body
                        email_subject = ReplacePlaceholdersEmailIndividual(Email_template.subject, replacements)
                        email_body = ReplacePlaceholdersEmailIndividual(Email_template.body, replacements)
                        Subject = email_subject
                        Body = email_body
                        #Builds an email message based on which recipients are included in the email
                        if to and cc and bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to, cc=cc, bcc=bcc)
                        elif to and cc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to, cc=cc)
                        elif to and bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to, bcc=bcc)
                        elif cc and bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, cc = cc, bcc=bcc)
                        elif to:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to)
                        elif cc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, cc = cc)
                        elif bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, bcc = bcc)
                        #Attaches the invoice pdf to the email
                        email.attach_file(invoice.invoice_file.path)
                        #Sends the email
                        email.send()
                #Reduces the number of invoice in invoice summary by one because one invoice was already sent (used for cron jobs tracking)
                invoice_summary.number_of_invoices -= 1
                invoice_summary.save()
            #Sends the json response to the front end
            return JsonResponse({'InvoicesToInvoiceItems': InvoicesTo_invoice_items})

        #Invoice handling if there is no contract that this invoice summary relates to
        else:
            if client_type == 'Client_company':
                Client_company = Client_Company.objects.get(id=client_id)
                #Creates the invoice summary object in the database
                invoice_summary = Invoice_Summary.objects.create(number_of_invoices = number_of_invoices, schedule=schedule, duration=duration, reminders_count=reminders_count, reminders_per=reminders_per, Total_price=Total_price, Total_VAT=Total_VAT, Total_with_VAT=Total_with_VAT, workplace=request.user.workplace, billed_to_company = Client_company)
                #Creates a unique identifier for the invoice (I thought that concatenating the company name with current date time down to seconds would guarantee uniqueness of invoice ID as well as serve as a good reference) (this is a personal choice)
                Company_name = (str(Client_company.name)).replace(' ', '')
                Current_date_time = timezone.now().strftime('%d%m%Y%H%M%S')
                identifier = (Company_name + Current_date_time).replace(" ", "")
                #Creates the invoice object in the database
                invoice = Invoice.objects.create(InvoiceSummary=invoice_summary, workplace = request.user.workplace, client_company=Client_company, identifier=identifier)
            else:
                Client_individual = Client_Individual.objects.get(id=client_id)
                #Creates the invoice summary object in the database
                invoice_summary = Invoice_Summary.objects.create(number_of_invoices = number_of_invoices, schedule=schedule, duration=duration, reminders_count=reminders_count, reminders_per=reminders_per, Total_price=Total_price, Total_VAT=Total_VAT, Total_with_VAT=Total_with_VAT, workplace=request.user.workplace, billed_to_individual = Client_individual)
                #Creates a unique identifier for the invoice (I thought that concatenating the company name with current date time down to seconds would guarantee uniqueness of invoice ID as well as serve as a good reference) (this is a personal choice)
                Individual_Fname = (str(Client_individual.Fname))
                print(Individual_Fname)
                Individual_Lname = (str(Client_individual.Lname))
                print(Individual_Lname)
                Current_date_time = timezone.now().strftime('%d%m%Y%H%M%S')
                print(Current_date_time)
                identifier = (Individual_Fname+Individual_Lname+Current_date_time).replace(" ", "")
                print(identifier)
                #Creates the invoice object in the database
                invoice = Invoice.objects.create(InvoiceSummary=invoice_summary, workplace = request.user.workplace, client_individual=Client_individual, identifier=identifier)
            #Extracts information about entities who the user configured to send emails with invoices to
            ContactsToBeTod = data.get("ContactsToBeTod")
            ContactsToBeCCd = data.get("ContactsToBeCCd")
            ContactsToBeBccd = data.get("ContactsToBeBccd")
            UsersToBeTod = data.get("UsersToBeTod")
            UsersToBeCCd = data.get("UsersToBeCCd")
            UsersToBeBccd = data.get("UsersToBeBccd")
            MainClientToBeTod = data.get("MainClientToBeTod")
            MainClientToBeCCd = data.get("MainClientToBeCCd")
            MainClientToBeBccd = data.get("MainClientToBeBccd")
            #Builds lists of entities to be included as different types of email recipients
            to = []
            cc = []
            bcc = []

            if client_type == 'Client_company':
                Client_company = Client_Company.objects.get(id=client_id)
                if MainClientToBeTod == True:
                    invoice_summary.MainClientToBeTod = True
                    invoice_summary.save()
                    to.append(Client_company.email)
                elif MainClientToBeCCd == True:
                    invoice_summary.MainClientToBeCCd = True
                    invoice_summary.save()
                    cc.append(Client_company.email)
                elif MainClientToBeBccd == True:
                    invoice_summary.MainClientToBeBccd = True
                    invoice_summary.save()
                    bcc.append(Client_company.email)
                else:
                    pass
            else:
                Client_individual = Client_Individual.objects.get(id=client_id)
                if MainClientToBeTod == True:
                    invoice_summary.MainClientToBeTod = True
                    invoice_summary.save()
                    to.append(Client_individual.email)
                elif MainClientToBeCCd == True:
                    invoice_summary.MainClientToBeCCd = True
                    invoice_summary.save()
                    cc.append(Client_individual.email)
                elif MainClientToBeBccd == True:
                    invoice_summary.MainClientToBeBccd = True
                    invoice_summary.save()
                    bcc.append(Client_individual.email)
                else:
                    pass

            if UsersToBeTod:
                for user in UsersToBeTod:
                    user = CustomUser.objects.get(id=user['id'], workplace=request.user.workplace)
                    UsersToBeTodThrough.objects.create(invoice_summary = invoice_summary, user = user, workplace = request.user.workplace)
                    to.append(user.email)
            if UsersToBeCCd:
                for user in UsersToBeCCd:
                    user = CustomUser.objects.get(id=user['id'], workplace=request.user.workplace)
                    UsersToBeCCdThrough.objects.create(invoice_summary = invoice_summary, user = user, workplace = request.user.workplace)
                    cc.append(user.email)
            if UsersToBeBccd:
                for user in UsersToBeBccd:
                    user = CustomUser.objects.get(id=user['id'], workplace=request.user.workplace)
                    UsersToBeBccdThrough.objects.create(invoice_summary = invoice_summary, user = user, workplace = request.user.workplace)
                    bcc.append(user.email)
            
            if ContactsToBeTod:
                for contact in ContactsToBeTod:
                    if contact['ClientType'] == 'Client_Company':
                        client_company = Client_Company.objects.get(id=contact['id'], workplace=request.user.workplace)
                        ContactsCompaniesToBeTodThrough.objects.create(invoice_summary=invoice_summary, client_company = client_company, workplace=request.user.workplace)
                        to.append(client_company.email)
                    else:
                        client_individual = Client_Individual.objects.get(id=contact['id'], workplace = request.user.workplace)
                        ContactsIndividualsToBeTodThrough.objects.create(invoice_summary=invoice_summary, client_individual=client_individual, workplace=request.user.workplace)
                        to.append(client_individual.email)
            if ContactsToBeCCd:
                for contact in ContactsToBeCCd:
                    if contact['ClientType'] == 'Client_Company':
                        client_company = Client_Company.objects.get(id=contact['id'], workplace=request.user.workplace)
                        ContactsCompaniesToBeCCdThrough.objects.create(invoice_summary=invoice_summary, client_company = client_company, workplace=request.user.workplace)
                        cc.append(client_company.email)
                    else:
                        client_individual = Client_Individual.objects.get(id=contact['id'], workplace = request.user.workplace)
                        ContactsIndividualsToBeCCdThrough.objects.create(invoice_summary=invoice_summary, client_individual=client_individual, workplace=request.user.workplace)
                        cc.append(client_individual.email)
            if ContactsToBeBccd:
                for contact in ContactsToBeBccd:
                    if contact['ClientType'] == 'Client_Company':
                        client_company = Client_Company.objects.get(id=contact['id'], workplace=request.user.workplace)
                        ContactsCompaniesToBeBccdThrough.objects.create(invoice_summary=invoice_summary, client_company = client_company, workplace=request.user.workplace)
                        bcc.append(client_company.email)
                    else:
                        client_individual = Client_Individual.objects.get(id=contact['id'], workplace = request.user.workplace)
                        ContactsIndividualsToBeBccdThrough.objects.create(invoice_summary=invoice_summary, client_individual=client_individual, workplace=request.user.workplace)
                        bcc.append(client_individual.email)
            #Adds entries of invoice items in the database and creates a list of them
            Invoice_schedule = data.get("Invoice_Schedule")
            ProductsOrServices = Invoice_schedule['products']
            invoice_items = []
            print('ProductsOrServices')
            for item in ProductsOrServices:
                Invoice_item = InvoiceItem.objects.create(Name=item['Name'], Price=item['PricePerInvoice'], Total_price=item['TotalPerInvoice'], VAT=item['VATPerInvoice'], invoice_summary=invoice_summary, workplace=request.user.workplace)
                invoice_items.append(Invoice_item)
                print(Invoice_item)
                #Because this handle invoice creation that doesn't relate to any contract the user can ditribute jobs when creating the invoice summary
                #This handles the job distribution of the invoice creation
                try:
                    Deadline = item['Deadline']
                    if Deadline != '':
                        Deadline = datetime.strptime(Deadline, "%Y-%m-%d").date()
                        if item['Employees']:
                            for user in item['Employees']:
                                job = Job.objects.create(corresponds_to_InvoiceItem=Invoice_item, workplace=request.user.workplace, Deadline=Deadline)
                                user = CustomUser.objects.get(id=user['id'], workplace=request.user.workplace)
                                jobassigned = JobAssigned.objects.create(Job=job, User=user)
                        else:
                            pass
                    else:
                        if item['Employees']:
                            for user in item['Employees']:
                                job = Job.objects.create(corresponds_to_InvoiceItem=Invoice_item, workplace=request.user.workplace)
                                user = CustomUser.objects.get(id=user['id'], workplace=request.user.workplace)
                                jobassigned = JobAssigned.objects.create(Job=job, User=user)
                        else:
                            pass
                except:
                    pass
            #Creates InvoicesToInvoiceItems object in the database to link the invoice with invoice items in the database
            InvoicesTo_invoice_items = []
            for invoice_item in invoice_items:
                invoices_to_invoice_items = InvoicesToInvoiceItems.objects.create(invoice_item = invoice_item, invoice=invoice, workplace=request.user.workplace)
                InvoicesTo_invoice_items.append(invoices_to_invoice_items.as_dict())
            
            if client_type == 'Client_company':
                print("Success")
                company_name = Client_company.name
                company_address = Client_company.address
                input_file = UploadedInvoiceWordDocumentCompany.objects.get(workplace=request.user.workplace)
                Date = datetime.today().strftime('%d-%m-%Y')
                #The due date of the invoice is set by me on the same day the invoice is created as this is a prototype application
                Due_date = datetime.today().strftime('%d-%m-%Y')
                Invoice_Number = invoice.identifier
                #Builds a table entries list to be used in the replace_placeholders_invoice function to populate the invoice with relevant data
                table_entries = []

                for item in ProductsOrServices:
                    entry = {'Name': item['Name'], 'Price': item['PricePerInvoice'], 'VAT': item['VATPerInvoice'], 'Total_price': item['TotalPerInvoice']}
                    table_entries.append(entry)
                print(table_entries)
                
                Total_price = invoice_summary.Total_price
                Total_VAT = invoice_summary.Total_VAT
                Total_with_VAT = invoice_summary.Total_with_VAT
                #Table total table entries list to be used to replace the total values in the invoice
                total_table_entries = {'Price': Total_price, 'VAT': Total_VAT, 'Price_with_VAT': Total_with_VAT}
                #Replacements dictionary to be used to replace placeholders with relevant information in the invoice using the replace_placeholders_invoice function
                replacements = {
                "{Company_Name}": company_name,
                "{Company_Address}": company_address,
                "{Date}": Date,
                "{Invoice_Number}": Invoice_Number,
                "{Due_date}": Due_date
                }
                #Creates the invoice using the replace_placeholders_invoice function
                replace_placeholders_invoice(input_file, replacements, table_entries, invoice, total_table_entries)
                #Creates a dictionary used for replacements for placeholders in the email that the invoice will be attached to
                if to or cc or bcc:
                        Email_template = EmailInvoiceWithoutContract_company.objects.get(workplace=request.user.workplace)
                        content_type = ContentType.objects.get_for_model(Client_company)
                        Related_Contact_First_Name = None
                        Related_Contact_Last_Name = None
                        Related_Contact_Company_Name = None
                        represents_list = Represents.objects.filter(from_content_type=content_type, from_object_id = Client_company.id, workplace=request.user.workplace, selected=True)
                        for rep in represents_list:
                            if isinstance(rep.to_entity, Client_Individual):
                                Related_Contact_First_Name = rep.to_entity.Fname
                                Related_Contact_Last_Name = rep.to_entity.Lname
                            elif isinstance(rep.to_entity, Client_Company):
                                Related_Contact_Company_Name = rep.to_entity.name
                        Signer = Client_company.Signer
                        Signer_First_Name = Signer.Fname
                        Signer_Last_Name = Signer.Lname
                        #The value of {Period} placeholder is chosen by me
                        if invoice_summary.schedule == "One Invoice":
                            current_date = time.strftime("%d-%m-%Y")
                            invoice_schedule = current_date
                        elif invoice_summary.schedule == "Daily":
                            current_date = time.strftime("%d-%m-%Y")
                            invoice_schedule = current_date
                        elif invoice_summary.schedule == "Weekly":
                            today = time.localtime()
                            week_number = int(time.strftime("%w", today))
                            weekday = int(time.strftime("%u", today))
                            seconds_in_day = 86400
                            today_timestamp = time.mktime(today)
                            days_since_monday = weekday - 1
                            monday_timestamp = today_timestamp - (days_since_monday * seconds_in_day)
                            monday = time.localtime(monday_timestamp)
                            today_str = time.strftime("%d-%m-%Y", today)
                            monday_str = time.strftime("%d-%m-%Y", monday)
                            invoice_schedule = f"week staring on {monday_str}"
                        elif invoice_summary.schedule == "Monthly":
                            current_month_date = time.strftime("%B-%Y")
                            invoice_schedule = current_month_date
                        elif invoice_summary.schedule == "Quarterly":
                            month = int(time.strftime("%m"))
                            quarter = (month - 1) // 3 + 1
                            current_year_date = time.strftime("%Y")
                            invoice_schedule = f"q{quarter} of {current_year_date}"
                        elif invoice_summary.schedule == "Yearly":
                            current_year_date = time.strftime("%Y")
                            invoice_schedule = current_year_date
                        replacements = {
                            "{Company_Name}": company_name,
                            "{Signer_First_Name}": Signer_First_Name,
                            "{Signer_Last_Name}": Signer_Last_Name,
                            "{Period}": invoice_schedule,
                            "{Related_Contact_First_Name}": Related_Contact_First_Name,
                            "{Related_Contact_Last_Name}": Related_Contact_Last_Name,
                            "{Related_Contact_Company_Name}": Related_Contact_Company_Name,
                        }
                        #Replaces the placeholders in the email template subject and body
                        email_subject = ReplacePlaceholdersEmailCompany(Email_template.subject, replacements)
                        email_body = ReplacePlaceholdersEmailCompany(Email_template.body, replacements)
                        Subject = email_subject
                        Body = email_body
                        #Builds an email message based on which recipients are included in the email
                        if to and cc and bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to, cc=cc, bcc=bcc)
                        elif to and cc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to, cc=cc)
                        elif to and bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to, bcc=bcc)
                        elif cc and bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, cc = cc, bcc=bcc)
                        elif to:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to)
                        elif cc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, cc = cc)
                        elif bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, bcc = bcc)
                        #Attaches the invoice pdf to the email
                        email.attach_file(invoice.invoice_file.path)
                        #Sends the email
                        email.send()
                #Reduces the number of invoice in invoice summary by one because one invoice was already sent (used for cron jobs tracking)
                invoice_summary.number_of_invoices -= 1
                invoice_summary.save()
            else:
                print("Success")
                individual_name = f"{Individual_Fname} {Individual_Lname}"
                individual_address = Client_individual.address
                input_file = UploadedInvoiceWordDocumentIndividual.objects.get(workplace=request.user.workplace)
                Date = datetime.today().strftime('%d-%m-%Y')
                #The due date of the invoice is set by me on the same day the invoice is created as this is a prototype application
                Due_date = datetime.today().strftime('%d-%m-%Y')
                Invoice_Number = invoice.identifier
                #Builds a table entries list to be used in the replace_placeholders_invoice function to populate the invoice with relevant data
                table_entries = []

                for item in ProductsOrServices:
                    entry = {'Name': item['Name'], 'Price': item['PricePerInvoice'], 'VAT': item['VATPerInvoice'], 'Total_price': item['TotalPerInvoice']}
                    table_entries.append(entry)
                print(table_entries)
                
                Total_price = invoice_summary.Total_price
                Total_VAT = invoice_summary.Total_VAT
                Total_with_VAT = invoice_summary.Total_with_VAT
                #Table total table entries list to be used to replace the total values in the invoice
                total_table_entries = {'Price': Total_price, 'VAT': Total_VAT, 'Price_with_VAT': Total_with_VAT}
                #Replacements dictionary to be used to replace placeholders with relevant information in the invoice using the replace_placeholders_invoice function
                replacements = {
                "{Client_First_Name}": Individual_Fname,
                "{Client_Last_Name}": Individual_Lname,
                "{Client_Address}": individual_address,
                "{Date}": Date,
                "{Invoice_Number}": Invoice_Number,
                "{Due_date}": Due_date
                }
                #Creates the invoice using the replace_placeholders_invoice function
                replace_placeholders_invoice(input_file, replacements, table_entries, invoice, total_table_entries)
                #Creates a dictionary used for replacements for placeholders in the email that the invoice will be attached to
                if to or cc or bcc:
                        Email_template = EmailInvoiceWithoutContract_individual.objects.get(workplace=request.user.workplace)
                        content_type = ContentType.objects.get_for_model(Client_individual)
                        Related_Contact_First_Name = None
                        Related_Contact_Last_Name = None
                        Related_Contact_Company_Name = None
                        represents_list = Represents.objects.filter(from_content_type=content_type, from_object_id = Client_individual.id, workplace=request.user.workplace, selected=True)
                        for rep in represents_list:
                            if isinstance(rep.to_entity, Client_Individual):
                                Related_Contact_First_Name = rep.to_entity.Fname
                                Related_Contact_Last_Name = rep.to_entity.Lname
                            elif isinstance(rep.to_entity, Client_Company):
                                Related_Contact_Company_Name = rep.to_entity.name
                        Signer = Client_individual
                        Signer_First_Name = Signer.Fname
                        Signer_Last_Name = Signer.Lname
                        #The value of {Period} placeholder is chosen by me
                        if invoice_summary.schedule == "One Invoice":
                            current_date = time.strftime("%d-%m-%Y")
                            invoice_schedule = current_date
                        elif invoice_summary.schedule == "Daily":
                            current_date = time.strftime("%d-%m-%Y")
                            invoice_schedule = current_date
                        elif invoice_summary.schedule == "Weekly":
                            today = time.localtime()
                            week_number = int(time.strftime("%w", today))
                            weekday = int(time.strftime("%u", today))
                            seconds_in_day = 86400
                            today_timestamp = time.mktime(today)
                            days_since_monday = weekday - 1
                            monday_timestamp = today_timestamp - (days_since_monday * seconds_in_day)
                            monday = time.localtime(monday_timestamp)
                            today_str = time.strftime("%d-%m-%Y", today)
                            monday_str = time.strftime("%d-%m-%Y", monday)
                            invoice_schedule = f"week staring on {monday_str}"
                        elif invoice_summary.schedule == "Monthly":
                            current_month_date = time.strftime("%B-%Y")
                            invoice_schedule = current_month_date
                        elif invoice_summary.schedule == "Quarterly":
                            month = int(time.strftime("%m"))
                            quarter = (month - 1) // 3 + 1
                            current_year_date = time.strftime("%Y")
                            invoice_schedule = f"q{quarter} of {current_year_date}"
                        elif invoice_summary.schedule == "Yearly":
                            current_year_date = time.strftime("%Y")
                            invoice_schedule = current_year_date
                        replacements = {
                            "{Client_First_Name}": Individual_Fname,
                            "{Client_Last_Name}": Individual_Lname,
                            "{Period}": invoice_schedule,
                            "{Related_Contact_First_Name}": Related_Contact_First_Name,
                            "{Related_Contact_Last_Name}": Related_Contact_Last_Name,
                            "{Related_Contact_Company_Name}": Related_Contact_Company_Name,
                        }
                        #Replaces the placeholders in the email template subject and body
                        email_subject = ReplacePlaceholdersEmailIndividual(Email_template.subject, replacements)
                        email_body = ReplacePlaceholdersEmailIndividual(Email_template.body, replacements)
                        Subject = email_subject
                        Body = email_body
                        #Builds an email message based on which recipients are included in the email
                        if to and cc and bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to, cc=cc, bcc=bcc)
                        elif to and cc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to, cc=cc)
                        elif to and bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to, bcc=bcc)
                        elif cc and bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, cc = cc, bcc=bcc)
                        elif to:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, to = to)
                        elif cc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, cc = cc)
                        elif bcc:
                            email = EmailMessage(subject = Subject, body = Body, from_email=None, bcc = bcc)
                        #Attaches the invoice pdf to the email
                        email.attach_file(invoice.invoice_file.path)
                        #Sends the email
                        email.send()
                #Reduces the number of invoice in invoice summary by one because one invoice was already sent (used for cron jobs tracking)
                invoice_summary.number_of_invoices -= 1
                invoice_summary.save()
            
            #Sends the json response to the front end
            return JsonResponse({'InvoicesToInvoiceItems': InvoicesTo_invoice_items})
        
    #Used to mark invoices as paid or not paid (in case of a mistake by the user) which will affect corn jobs for sending email reminders for this invoice
    elif request.method == 'PUT':   
        data = json.loads(request.body) #Retrieves the request body and turns it into python readable dictionary
        invoice_id = data.get('Invoice')['id'] #Extracts the invoice id from the request body dictionary
        paid = data.get('Paid') #Extracts the Paid value from the request body dictionary
        #Retrieves the invoice from the database
        invoice = Invoice.objects.get(id=invoice_id, workplace=request.user.workplace)
        #Sets the invoice.Paid value the extracted value from the request body dictionary
        invoice.Paid = paid
        invoice.save()
        #If this invoice relates to a contract, accumulates the value of the invoice to paid value of the contract
        if invoice.contract and paid == True:
            contract = invoice.contract
            contract.Paid += invoice.InvoiceSummary.Total_with_VAT
            contract.save()
            return JsonResponse({'invoice': invoice.as_dict(), 'contract': contract.as_dict()})
        #If this invoice relates to a contract, subtracts the value of the invoice from the paid value of the contract
        elif invoice.contract and paid == False:
            contract = invoice.contract
            contract.Paid -= invoice.InvoiceSummary.Total_with_VAT
            contract.save()
        #Returns a response to the front end
            return JsonResponse({'invoice': invoice.as_dict(), 'contract': contract.as_dict()})
        else:
            return JsonResponse({'invoice': invoice.as_dict()})
    #Used to delete invoices which will also terminate all automations related to this invoice by deleting the invoice summary it is related to (Like this because it is a prototype web app)
    elif request.method == 'DELETE':
        data = json.loads(request.body)
        invoice_id = data.get('invoice_id')
        invoice = Invoice.objects.get(id = invoice_id, workplace = request.user.workplace)
        invoice_summary = invoice.InvoiceSummary
        invoice.delete()
        invoice_summary.delete()
        return JsonResponse({'DELETE': 'Confirm'})







    

#API endpoint for handling contract items
@csrf_exempt
@login_required
def api_ContractItem(request):
    if request.method == 'GET':
        data = json.loads(request.body)
        id = data.get('id')
        contract = data.get('contract')
        is_in = data.get('invoice')
        contract_item_get = ContractItem.objects.get(id=id, contract=Contract.objects.get(id=contract), is_in=Invoice.objects.get(id=is_in), workplace=request.user.workplace)
        return JsonResponse(contract_item_get.as_dict())
    
    elif request.method == 'POST':
        data = json.loads(request.body)
        ItemName = data.get('ItemName')
        description = data.get('description')
        amount = data.get('amount')
        contract = data.get('contract')
        is_in = data.get('invoice')
        contract_item_post = ContractItem.objects.create(ItemName = ItemName, description=description, amount=amount, contract=contract, is_in=is_in, workplace=request.user.workplace)
        return JsonResponse(contract_item_post.as_dict())
    
    elif request.method == 'PUT':
        data=json.loads(request.body)
        id = data.get('id')
        ItemName = data.get('ItemName')
        description = data.get('description')
        amount = data.get('amount')
        contract = data.get('contract')
        is_in = data.get('invoice')
        contract_item_put = ContractItem.objects.get(id=id, contract=contract, is_in=is_in, workplace=request.user.workplace)
        contract_item_put.ItemName = ItemName
        contract_item_put.description = description
        contract_item_put.amount = amount
        contract_item_put.save()
        return JsonResponse(contract_item_put.as_dict())
    
    elif request.method == 'DELETE':
        data = json.loads(request.body)
        id = data.get('id')
        contract = data.get('contract')
        is_in = data.get('invoice')
        contract_item_delete = ContractItem.objects.get(id=id, contract=contract, is_in=is_in, workplace = request.user.workplace)
        contract_item_delete.delete()
        return JsonResponse({'Contract Item': 'Deleted'})
    
    else:
        return HttpResponse(status=405)

#This function is used to handle requests associated with jobs the user has
@csrf_exempt
@login_required
def api_Job(request):
    #Returns a list of jobs the user has
    if request.method == 'GET':
        assigned_jobs = JobAssigned.objects.filter(User=request.user)
        return JsonResponse({
            'Jobs': [
                assignment.Job.as_dict() for assignment in assigned_jobs
            ],
        })
    #Used to mark jobs as started, not started, completed, not completed and record dates of start and completion
    elif request.method == 'PUT':
        data = json.loads(request.body)
        id = data.get('Job')['id']
        status = data.get('Job')['status']
        try:
            Started = data.get('Started')
        except:
            pass
        try:
            Finished = data.get('Finished')
        except:
            pass
        job_item_put = Job.objects.get(id=id, workplace=request.user.workplace)
        try:
            if Started == True and job_item_put.started_at is None:
                job_item_put.started_at = timezone.now()
            elif Started == False and job_item_put.started_at is not None:
                job_item_put.started_at = None
            else:
                pass
            if Finished == True and job_item_put.completed_at is None:
                job_item_put.completed_at = timezone.now()
                if job_item_put.corresponds_to_ContractItem:
                    Contractitem = job_item_put.corresponds_to_ContractItem
                    Contract = Contractitem.contract
                    Contract.Completed_Value += Contractitem.Total_price
                    invoices = Invoice.objects.filter(contract=Contract, workplace = request.user.workplace)
                    for invoice in invoices:
                        invoice_items = InvoicesToInvoiceItems.objects.filter(invoice = invoice)
                        for entry in invoice_items:
                            item = entry.invoice_item
                            if item.Name == Contractitem.Name:
                                invoice.Completed_Value += item.Total_price
                        invoice.save()
                    Contract.save()
                else:
                    Invoiceitem = job_item_put.corresponds_to_InvoiceItem
                    Invoices_to_invoice_items = InvoicesToInvoiceItems.objects.get(invoice_item = Invoiceitem)
                    invoice = Invoices_to_invoice_items.invoice
                    invoice.Completed_Value += Invoiceitem.Total_price
                    invoice.save()
            elif Finished == False and job_item_put.completed_at  is not None:
                job_item_put.completed_at = None
                if job_item_put.corresponds_to_ContractItem:
                    Contractitem = job_item_put.corresponds_to_ContractItem
                    Contract = Contractitem.contract
                    Contract.Completed_Value -= Contractitem.Total_price
                    invoices = Invoice.objects.filter(contract=Contract, workplace = request.user.workplace)
                    for invoice in invoices:
                        invoice_items = InvoicesToInvoiceItems.objects.filter(invoice = invoice)
                        for entry in invoice_items:
                            item = entry.invoice_item
                            if item.Name == Contractitem.Name:
                                invoice.Completed_Value -= item.Total_price
                        invoice.save()
                    Contract.save()
                else:
                    Invoiceitem = job_item_put.corresponds_to_InvoiceItem
                    Invoices_to_invoice_items = InvoicesToInvoiceItems.objects.get(invoice_item = Invoiceitem)
                    invoice = Invoices_to_invoice_items.invoice
                    invoice.Completed_Value -= Invoiceitem.Total_price
                    invoice.save()
            else:
                pass
        except:
            pass

        if status != job_item_put.status:
            job_item_put.status = status
        else:
            pass

        job_item_put.save()

        return JsonResponse(job_item_put.as_dict())
    #Used to delete jobs, not used in the prototype.
    elif request.method == 'DELETE':
        data = json.loads(request.body)
        id = data.get('id')
        corresponds_to = data.get('ContractItem')
        job_item_delete = Job.objects.get(id=id, corresponds_to=corresponds_to, user = request.user, workplace=request.user.workplace)
        job_item_delete.delete()
        return JsonResponse({'Job item': 'deleted'})
    
    else:
        return HttpResponse(status=405)
    

